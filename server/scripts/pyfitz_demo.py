import os
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import cv2
import numpy as np
from pathlib import Path
from typing import Dict, Optional, List, Tuple
import io

# --- DOCX Text Extractor ---
class DocxTextExtractor:
    """
    DOCX text extraction pipeline that extracts text and images with OCR using EasyOCR.
    Similar structure to PDFTextExtractor but uses EasyOCR instead of pytesseract.
    """
    
    def __init__(self, enable_gpu: bool = True, output_dir: str = "extracted_texts"):
        """Initialize the DOCX text extractor with EasyOCR."""
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # Setup logging FIRST (before using logger)
        import logging
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger(__name__)
        
        # Initialize EasyOCR extractor
        try:
            self.ocr_extractor = EasyOCRExtractor(enable_gpu=enable_gpu)
            self.logger.info("EasyOCR initialized successfully for DOCX processing")
        except Exception as e:
            self.logger.error(f"Failed to initialize EasyOCR: {e}")
            raise
    
    def diagnose_docx(self, docx_path: str) -> Dict[str, any]:
        """
        Enhanced DOCX structure diagnosis with detailed image and document metadata analysis.
        Similar to PDFTextExtractor's comprehensive diagnostics.
        """
        try:
            from docx import Document
            import zipfile
            import xml.etree.ElementTree as ET
            
            doc = Document(docx_path)
            
            # Basic document structure
            diagnosis = {
                'docx_path': docx_path,
                'file_size_bytes': os.path.getsize(docx_path),
                'total_paragraphs': len(doc.paragraphs),
                'total_tables': len(doc.tables),
                'total_sections': len(doc.sections),
                'images_found': 0,
                'image_details': [],
                'content_analysis': {},
                'document_properties': {},
                'media_files': []
            }
            
            # Document properties and metadata
            try:
                core_props = doc.core_properties
                diagnosis['document_properties'] = {
                    'title': getattr(core_props, 'title', '') or '',
                    'author': getattr(core_props, 'author', '') or '',
                    'subject': getattr(core_props, 'subject', '') or '',
                    'created': str(getattr(core_props, 'created', '')) if getattr(core_props, 'created', None) else '',
                    'modified': str(getattr(core_props, 'modified', '')) if getattr(core_props, 'modified', None) else '',
                    'last_modified_by': getattr(core_props, 'last_modified_by', '') or '',
                    'revision': getattr(core_props, 'revision', 0) or 0
                }
            except Exception as e:
                self.logger.debug(f"Could not extract document properties: {e}")
                diagnosis['document_properties'] = {'error': str(e)}
            
            # Content analysis
            total_text_length = 0
            non_empty_paragraphs = 0
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    non_empty_paragraphs += 1
                    total_text_length += len(text)
            
            # Table analysis
            table_analysis = []
            for i, table in enumerate(doc.tables):
                table_info = {
                    'table_index': i + 1,
                    'rows': len(table.rows),
                    'columns': len(table.columns) if table.rows else 0,
                    'total_cells': len(table.rows) * (len(table.columns) if table.rows else 0)
                }
                table_analysis.append(table_info)
            
            diagnosis['content_analysis'] = {
                'total_text_length': total_text_length,
                'non_empty_paragraphs': non_empty_paragraphs,
                'empty_paragraphs': len(doc.paragraphs) - non_empty_paragraphs,
                'average_paragraph_length': total_text_length / non_empty_paragraphs if non_empty_paragraphs > 0 else 0,
                'table_analysis': table_analysis
            }
            
            # Detailed image analysis (similar to PDFTextExtractor)
            with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                # Find all media files
                media_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
                diagnosis['media_files'] = media_files
                
                # Analyze images in detail
                image_files = [f for f in media_files 
                             if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff'))]
                diagnosis['images_found'] = len(image_files)
                
                for img_idx, img_file in enumerate(image_files):
                    try:
                        # Extract image data for analysis
                        image_bytes = docx_zip.read(img_file)
                        
                        # Get basic image info
                        img_detail = {
                            'index': img_idx + 1,
                            'filename': os.path.basename(img_file),
                            'path_in_docx': img_file,
                            'size_bytes': len(image_bytes),
                            'format': 'unknown',
                            'width': 0,
                            'height': 0,
                            'colorspace': 'unknown',
                            'mode': 'unknown'
                        }
                        
                        # Analyze image using PIL
                        try:
                            from PIL import Image
                            import io
                            
                            image = Image.open(io.BytesIO(image_bytes))
                            img_detail.update({
                                'format': image.format or 'unknown',
                                'width': image.width,
                                'height': image.height,
                                'mode': image.mode,
                                'colorspace': image.mode,
                                'has_transparency': 'transparency' in image.info or image.mode in ('RGBA', 'LA'),
                                'total_pixels': image.width * image.height
                            })
                            
                            # Calculate compression ratio
                            expected_size = img_detail['total_pixels'] * (4 if image.mode == 'RGBA' else 3)
                            img_detail['compression_ratio'] = expected_size / len(image_bytes) if len(image_bytes) > 0 else 0
                            
                        except Exception as e:
                            img_detail['image_analysis_error'] = str(e)
                            self.logger.debug(f"Could not analyze image {img_file}: {e}")
                        
                        diagnosis['image_details'].append(img_detail)
                        
                    except Exception as e:
                        diagnosis['image_details'].append({
                            'index': img_idx + 1,
                            'filename': os.path.basename(img_file),
                            'error': str(e)
                        })
            
            # Summary statistics
            diagnosis['summary'] = {
                'total_media_files': len(diagnosis['media_files']),
                'total_image_bytes': sum(img.get('size_bytes', 0) for img in diagnosis['image_details']),
                'average_image_size': sum(img.get('size_bytes', 0) for img in diagnosis['image_details']) / len(diagnosis['image_details']) if diagnosis['image_details'] else 0,
                'largest_image_bytes': max((img.get('size_bytes', 0) for img in diagnosis['image_details']), default=0),
                'total_pixels': sum(img.get('total_pixels', 0) for img in diagnosis['image_details']),
                'document_complexity_score': self._calculate_complexity_score(diagnosis)
            }
            
            self.logger.info(f"Enhanced DOCX diagnosis completed: {diagnosis['total_paragraphs']} paragraphs, "
                           f"{diagnosis['total_tables']} tables, {diagnosis['images_found']} images, "
                           f"complexity score: {diagnosis['summary']['document_complexity_score']:.1f}")
            
            return diagnosis
            
        except Exception as e:
            self.logger.error(f"Enhanced DOCX diagnosis failed: {e}")
            return {'error': str(e)}
    
    def _calculate_complexity_score(self, diagnosis: Dict) -> float:
        """
        Calculate a complexity score for the document based on various factors.
        Similar to how PDFTextExtractor analyzes document complexity.
        """
        try:
            score = 0.0
            
            # Text complexity (30% weight)
            content = diagnosis.get('content_analysis', {})
            text_length = content.get('total_text_length', 0)
            paragraphs = content.get('non_empty_paragraphs', 0)
            score += min(text_length / 10000, 1.0) * 30  # Max 30 points for text
            score += min(paragraphs / 50, 1.0) * 20  # Max 20 points for paragraph count
            
            # Table complexity (20% weight)
            tables = content.get('table_analysis', [])
            if tables:
                total_cells = sum(t.get('total_cells', 0) for t in tables)
                score += min(total_cells / 100, 1.0) * 20  # Max 20 points for tables
            
            # Image complexity (30% weight)
            images = diagnosis.get('image_details', [])
            if images:
                total_pixels = sum(img.get('total_pixels', 0) for img in images)
                image_count = len(images)
                score += min(total_pixels / 10000000, 1.0) * 20  # Max 20 points for pixel count
                score += min(image_count / 10, 1.0) * 10  # Max 10 points for image count
            
            # Document structure complexity (20% weight)
            sections = diagnosis.get('total_sections', 0)
            media_files = len(diagnosis.get('media_files', []))
            score += min(sections / 5, 1.0) * 10  # Max 10 points for sections
            score += min(media_files / 20, 1.0) * 10  # Max 10 points for media files
            
            return min(score, 100.0)  # Cap at 100
            
        except Exception as e:
            self.logger.debug(f"Complexity score calculation failed: {e}")
            return 0.0
    
    def extract_native_text(self, docx_path: str) -> Dict:
        """Extract native text from DOCX using python-docx."""
        try:
            from docx import Document
            doc = Document(docx_path)
            
            # Extract paragraphs
            paragraphs = []
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    paragraphs.append({
                        'paragraph_num': i + 1,
                        'text': paragraph.text.strip()
                    })
            
            # Extract tables
            tables = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                tables.append({
                    'table_num': table_idx + 1,
                    'data': table_data,
                    'rows': len(table_data),
                    'cols': len(table_data[0]) if table_data else 0
                })
            
            # Extract headers and footers
            headers = []
            footers = []
            for section_idx, section in enumerate(doc.sections):
                if section.header:
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            headers.append(para.text.strip())
                if section.footer:
                    for para in section.footer.paragraphs:
                        if para.text.strip():
                            footers.append(para.text.strip())
            
            text_data = {
                'paragraphs': paragraphs,
                'tables': tables,
                'headers': headers,
                'footers': footers,
                'total_paragraphs': len(paragraphs),
                'total_tables': len(tables)
            }
            
            self.logger.info(f"Extracted {len(paragraphs)} paragraphs and {len(tables)} tables from DOCX")
            return text_data
            
        except ImportError:
            self.logger.error("python-docx not installed. Install with: pip install python-docx")
            return {'error': 'python-docx not installed'}
        except Exception as e:
            self.logger.error(f"Failed to extract native text: {e}")
            # Fallback to docx2txt if available
            try:
                import docx2txt
                text = docx2txt.process(docx_path)
                return {
                    'paragraphs': [{'paragraph_num': 1, 'text': text}],
                    'tables': [],
                    'headers': [],
                    'footers': [],
                    'fallback_used': True
                }
            except:
                return {'error': str(e)}
    
    def extract_images_from_docx(self, docx_path: str) -> List[Dict]:
        """Extract images from DOCX file and save them to extracted_images folder."""
        images = []
        
        # Create extracted_images folder
        extracted_images_dir = self.output_dir / "extracted_images"
        extracted_images_dir.mkdir(exist_ok=True)
        
        try:
            import zipfile
            
            # DOCX is essentially a ZIP file
            with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                # Find all image files in the media folder
                image_files = [f for f in docx_zip.namelist() 
                             if f.startswith('word/media/') and 
                             f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff'))]
                
                self.logger.info(f"Found {len(image_files)} images in DOCX")
                
                for img_index, img_file in enumerate(image_files):
                    try:
                        # Extract image data
                        image_bytes = docx_zip.read(img_file)
                        
                        # Filter out tiny images (likely decorative elements)
                        if len(image_bytes) < 1000:
                            self.logger.debug(f"Skipping tiny image {img_index + 1} (size: {len(image_bytes)} bytes)")
                            continue
                        
                        # Convert to PIL Image
                        image = Image.open(io.BytesIO(image_bytes))
                        
                        # Filter out very small images by dimensions
                        if image.width < 50 or image.height < 50:
                            self.logger.debug(f"Skipping small image {img_index + 1} (dimensions: {image.width}x{image.height})")
                            continue
                        
                        # Save the image to disk
                        original_name = Path(img_file).name
                        img_filename = f"docx_img{img_index + 1}_{original_name}"
                        img_path = extracted_images_dir / img_filename
                        image.save(img_path)
                        
                        image_data = {
                            "image": image,
                            "index": img_index + 1,
                            "size": (image.width, image.height),
                            "type": "image",
                            "saved_path": str(img_path),
                            "filename": img_filename,
                            "original_name": original_name
                        }
                        
                        images.append(image_data)
                        self.logger.info(f"Successfully extracted and saved image {img_index + 1} (size: {image.width}x{image.height}, {len(image_bytes)} bytes) -> {img_filename}")
                        
                    except Exception as e:
                        self.logger.warning(f"Failed to extract image {img_index + 1} from DOCX: {e}")
            
            self.logger.info(f"Successfully extracted {len(images)} valid images from DOCX")
            return images
            
        except Exception as e:
            self.logger.error(f"Failed to extract images from DOCX: {e}")
            return []
    
    def advanced_preprocess_image(self, image_path: str) -> List[Tuple[str, np.ndarray]]:
        """
        Apply multiple advanced preprocessing methods similar to PDFTextExtractor.
        Returns list of (method_name, processed_image) tuples.
        """
        try:
            # Load image
            image = cv2.imread(image_path)
            if image is None:
                raise ValueError(f"Could not load image: {image_path}")
            
            processed_images = []
            
            # Convert to grayscale
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Method 1: EasyOCR's original preprocessing (from existing EasyOCRExtractor)
            try:
                easyocr_processed = self.ocr_extractor.preprocess_image(image_path)
                processed_images.append(('easyocr_enhanced', easyocr_processed))
                self.logger.debug("Applied EasyOCR preprocessing")
            except Exception as e:
                self.logger.debug(f"EasyOCR preprocessing failed: {e}")
            
            # Method 2: Adaptive thresholding (from PDFTextExtractor)
            try:
                # Denoise first
                denoised = cv2.bilateralFilter(gray, 9, 75, 75)
                adaptive_thresh = cv2.adaptiveThreshold(
                    denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
                )
                processed_images.append(('adaptive_threshold', adaptive_thresh))
                self.logger.debug("Applied adaptive thresholding")
            except Exception as e:
                self.logger.debug(f"Adaptive thresholding failed: {e}")
            
            # Method 3: OTSU thresholding
            try:
                _, otsu_thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                processed_images.append(('otsu_threshold', otsu_thresh))
                self.logger.debug("Applied OTSU thresholding")
            except Exception as e:
                self.logger.debug(f"OTSU thresholding failed: {e}")
            
            # Method 4: Enhanced contrast with CLAHE + thresholding
            try:
                clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
                enhanced = clahe.apply(gray)
                _, clahe_thresh = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                processed_images.append(('clahe_enhanced', clahe_thresh))
                self.logger.debug("Applied CLAHE enhancement")
            except Exception as e:
                self.logger.debug(f"CLAHE enhancement failed: {e}")
            
            # Method 5: Morphological operations
            try:
                kernel = np.ones((2,2), np.uint8)
                morph = cv2.morphologyEx(gray, cv2.MORPH_CLOSE, kernel)
                _, morph_thresh = cv2.threshold(morph, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                processed_images.append(('morphological', morph_thresh))
                self.logger.debug("Applied morphological operations")
            except Exception as e:
                self.logger.debug(f"Morphological operations failed: {e}")
            
            # Method 6: Simple thresholding
            try:
                _, simple_thresh = cv2.threshold(gray, 127, 255, cv2.THRESH_BINARY)
                processed_images.append(('simple_threshold', simple_thresh))
                self.logger.debug("Applied simple thresholding")
            except Exception as e:
                self.logger.debug(f"Simple thresholding failed: {e}")
            
            # Method 7: Scaled image (2x resolution)
            try:
                height, width = gray.shape[:2]
                scaled = cv2.resize(gray, (width * 2, height * 2), interpolation=cv2.INTER_CUBIC)
                processed_images.append(('scaled_2x', scaled))
                self.logger.debug("Applied 2x scaling")
            except Exception as e:
                self.logger.debug(f"Image scaling failed: {e}")
            
            self.logger.info(f"Generated {len(processed_images)} preprocessed versions of {Path(image_path).name}")
            return processed_images
            
        except Exception as e:
            self.logger.error(f"Advanced preprocessing failed for {image_path}: {e}")
            return []

    def extract_text_from_image(self, image_path: str) -> Dict:
        """
        Extract text from saved image using EasyOCR with comprehensive fallback strategies.
        Implements multiple OCR attempts with different EasyOCR configurations and preprocessing.
        Returns structured output with text, confidence, and word count.
        """
        try:
            self.logger.info(f"Starting comprehensive OCR pipeline for {Path(image_path).name}")
            
            all_ocr_results = []
            
            # Phase 1: Try original image with different EasyOCR configurations
            original_configs = [
                ('original_standard', {'gpu': True, 'paragraph': False}),
                ('original_paragraph', {'gpu': True, 'paragraph': True}),
                ('original_no_gpu', {'gpu': False, 'paragraph': False})
            ]
            
            for config_name, config_params in original_configs:
                try:
                    # Create temporary EasyOCR reader with different config
                    import easyocr
                    temp_reader = easyocr.Reader(['en'], **config_params)
                    
                    # Extract text with this configuration
                    results = temp_reader.readtext(image_path)
                    texts = []
                    confidences = []
                    
                    for (bbox, text, confidence) in results:
                        if confidence > 0.1:  # Filter low confidence
                            texts.append(text.strip())
                            confidences.append(confidence)
                    
                    avg_confidence = sum(confidences) / len(confidences) if confidences else 0
                    result = {
                        'text': '\n'.join(texts),
                        'confidence': avg_confidence,
                        'word_count': len(texts),
                        'method': config_name,
                        'preprocessing_applied': False,
                        'config_used': config_params
                    }
                    
                    all_ocr_results.append(result)
                    self.logger.debug(f"{config_name} OCR: {len(result['text'])} chars, confidence: {avg_confidence:.2f}")
                    
                except Exception as e:
                    self.logger.debug(f"EasyOCR config {config_name} failed: {e}")
                    # Add a fallback result with error info
                    all_ocr_results.append({
                        'text': '',
                        'confidence': 0,
                        'word_count': 0,
                        'method': config_name,
                        'preprocessing_applied': False,
                        'error': str(e)
                    })
            
            # Phase 2: Try preprocessed images with standard EasyOCR
            preprocessed_images = self.advanced_preprocess_image(image_path)
            
            for method_name, processed_image in preprocessed_images:
                try:
                    # Save preprocessed image temporarily
                    import tempfile
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                        temp_path = temp_file.name
                        cv2.imwrite(temp_path, processed_image)
                    
                    try:
                        # Apply standard EasyOCR to preprocessed image
                        result = self.ocr_extractor.extract_text(temp_path)
                        result['method'] = method_name
                        result['preprocessing_applied'] = True
                        all_ocr_results.append(result)
                        
                        self.logger.debug(f"{method_name} OCR: {len(result.get('text', ''))} chars, confidence: {result.get('confidence', 0):.2f}")
                        
                    finally:
                        # Clean up temporary file
                        try:
                            os.unlink(temp_path)
                        except:
                            pass
                            
                except Exception as e:
                    self.logger.debug(f"OCR with {method_name} preprocessing failed: {e}")
            
            # Phase 3: Advanced fallback - try different image manipulations
            advanced_fallbacks = self._generate_advanced_fallback_images(image_path)
            
            for fallback_name, fallback_image_path in advanced_fallbacks:
                try:
                    result = self.ocr_extractor.extract_text(fallback_image_path)
                    result['method'] = f'fallback_{fallback_name}'
                    result['preprocessing_applied'] = True
                    result['is_fallback'] = True
                    all_ocr_results.append(result)
                    
                    self.logger.debug(f"Fallback {fallback_name} OCR: {len(result.get('text', ''))} chars, confidence: {result.get('confidence', 0):.2f}")
                    
                    # Clean up fallback image
                    try:
                        os.unlink(fallback_image_path)
                    except:
                        pass
                        
                except Exception as e:
                    self.logger.debug(f"Fallback {fallback_name} OCR failed: {e}")
            
            # Phase 4: Select the best result using enhanced scoring
            if all_ocr_results:
                best_result = self._select_best_ocr_result(all_ocr_results)
                
                # Add comprehensive metadata
                best_result['total_attempts'] = len(all_ocr_results)
                best_result['all_methods'] = [r['method'] for r in all_ocr_results]
                best_result['score_distribution'] = [self._score_ocr_result(r) for r in all_ocr_results]
                
                self.logger.info(f"Best OCR result from '{best_result['method']}': "
                               f"{len(best_result.get('text', ''))} chars, "
                               f"confidence: {best_result.get('confidence', 0):.2f}, "
                               f"score: {best_result.get('best_score', 0):.1f} "
                               f"(tried {len(all_ocr_results)} methods)")
                
                return best_result
            else:
                self.logger.warning("All comprehensive OCR methods failed")
                return {'text': '', 'confidence': 0, 'word_count': 0, 'error': 'All OCR methods failed', 'preprocessing_applied': False}
                
        except Exception as e:
            self.logger.error(f"Comprehensive OCR extraction failed for {image_path}: {e}")
            return {'text': '', 'confidence': 0, 'word_count': 0, 'error': str(e), 'preprocessing_applied': False}
    
    def _generate_advanced_fallback_images(self, image_path: str) -> List[Tuple[str, str]]:
        """
        Generate advanced fallback images with extreme preprocessing for difficult cases.
        Returns list of (method_name, temp_image_path) tuples.
        """
        fallback_images = []
        
        try:
            import tempfile
            
            # Load original image
            image = cv2.imread(image_path)
            if image is None:
                return fallback_images
            
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Fallback 1: Extreme contrast enhancement
            try:
                clahe = cv2.createCLAHE(clipLimit=4.0, tileGridSize=(4,4))
                extreme_contrast = clahe.apply(gray)
                
                temp_file = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
                cv2.imwrite(temp_file.name, extreme_contrast)
                fallback_images.append(('extreme_contrast', temp_file.name))
                temp_file.close()
            except Exception as e:
                self.logger.debug(f"Extreme contrast fallback failed: {e}")
            
            # Fallback 2: Multiple scale attempts
            for scale_factor in [3, 4, 0.5]:
                try:
                    height, width = gray.shape[:2]
                    new_width = int(width * scale_factor)
                    new_height = int(height * scale_factor)
                    
                    if new_width > 0 and new_height > 0 and new_width < 10000 and new_height < 10000:
                        scaled = cv2.resize(gray, (new_width, new_height), interpolation=cv2.INTER_CUBIC)
                        
                        temp_file = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
                        cv2.imwrite(temp_file.name, scaled)
                        fallback_images.append((f'scale_{scale_factor}x', temp_file.name))
                        temp_file.close()
                except Exception as e:
                    self.logger.debug(f"Scale {scale_factor}x fallback failed: {e}")
            
            # Fallback 3: Edge enhancement
            try:
                # Sharpen the image
                kernel = np.array([[-1,-1,-1], [-1,9,-1], [-1,-1,-1]])
                sharpened = cv2.filter2D(gray, -1, kernel)
                
                temp_file = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
                cv2.imwrite(temp_file.name, sharpened)
                fallback_images.append(('edge_enhanced', temp_file.name))
                temp_file.close()
            except Exception as e:
                self.logger.debug(f"Edge enhancement fallback failed: {e}")
            
            # Fallback 4: Gaussian blur + threshold (for noisy images)
            try:
                blurred = cv2.GaussianBlur(gray, (3, 3), 0)
                _, thresh = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                
                temp_file = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
                cv2.imwrite(temp_file.name, thresh)
                fallback_images.append(('blur_threshold', temp_file.name))
                temp_file.close()
            except Exception as e:
                self.logger.debug(f"Blur threshold fallback failed: {e}")
            
            self.logger.debug(f"Generated {len(fallback_images)} advanced fallback images")
            
        except Exception as e:
            self.logger.error(f"Failed to generate fallback images: {e}")
        
        return fallback_images
    
    def _score_ocr_result(self, result: Dict) -> float:
        """Enhanced scoring function for OCR results with multiple criteria."""
        try:
            text = result.get('text', '').strip()
            confidence = result.get('confidence', 0)
            word_count = result.get('word_count', 0)
            
            # Base score components
            text_length_score = min(len(text) / 1000, 1.0) * 40  # Max 40 points
            confidence_score = confidence * 30  # Max 30 points
            word_count_score = min(word_count / 50, 1.0) * 20  # Max 20 points
            
            # Bonus points for quality indicators
            bonus_score = 0
            
            # Bonus for reasonable text patterns
            if text:
                # Check for alphanumeric content
                alphanumeric_ratio = sum(1 for c in text if c.isalnum()) / len(text)
                bonus_score += alphanumeric_ratio * 5
                
                # Bonus for multiple words
                if word_count > 1:
                    bonus_score += 3
                
                # Bonus for common word patterns
                import re
                if re.search(r'\b(the|and|or|in|on|at|to|for|of|with|by)\b', text.lower()):
                    bonus_score += 2
            
            total_score = text_length_score + confidence_score + word_count_score + bonus_score
            return min(total_score, 100.0)  # Cap at 100
            
        except Exception as e:
            self.logger.debug(f"OCR scoring failed: {e}")
            return 0.0
    
    def _select_best_ocr_result(self, results: List[Dict]) -> Dict:
        """
        Select the best OCR result using enhanced multi-criteria scoring.
        """
        if not results:
            return {'text': '', 'confidence': 0, 'word_count': 0, 'error': 'No results to select from'}
        
        # Score all results
        scored_results = []
        for result in results:
            score = self._score_ocr_result(result)
            result['best_score'] = score
            scored_results.append((score, result))
        
        # Sort by score (descending)
        scored_results.sort(key=lambda x: x[0], reverse=True)
        
        # Additional validation: prefer results with actual text content
        for score, result in scored_results:
            if result.get('text', '').strip():  # Has meaningful text
                return result
        
        # If no result has text, return the highest scored one
        return scored_results[0][1] if scored_results else results[0]
    
    def process_docx(self, docx_path: str, diagnose: bool = True) -> Dict[str, any]:
        """Main method to process a DOCX file and extract all text content."""
        self.logger.info(f"Starting DOCX processing: {docx_path}")
        
        results = {
            'docx_path': docx_path,
            'native_text': {},
            'images': [],
            'image_ocr_text': [],
            'combined_text': '',
            'diagnosis': None
        }
        
        try:
            # Run diagnostics first if requested
            if diagnose:
                self.logger.info("Running DOCX diagnostics...")
                diagnosis = self.diagnose_docx(docx_path)
                results['diagnosis'] = diagnosis
            
            # Extract native text
            self.logger.info("Extracting native text...")
            native_text = self.extract_native_text(docx_path)
            results['native_text'] = native_text
            
            # Extract images and perform OCR
            self.logger.info("Extracting images and performing OCR...")
            images = self.extract_images_from_docx(docx_path)
            results['images'] = images
            
            # Process images with EasyOCR
            image_ocr_results = []
            for img_data in images:
                try:
                    image_path = img_data["saved_path"]  # Use the saved image file path
                    self.logger.info(f"Starting EasyOCR for image {img_data['index']} (size: {img_data['size']})")
                    
                    # Apply EasyOCR using the saved image file
                    ocr_result = self.extract_text_from_image(image_path)
                    
                    # Store the complete OCR result
                    img_data["ocr_result"] = ocr_result
                    img_data["ocr_text"] = ocr_result.get('text', '')
                    img_data["ocr_confidence"] = ocr_result.get('confidence', 0)
                    img_data["ocr_word_count"] = ocr_result.get('word_count', 0)
                    
                    # Create structured result for summary with advanced OCR info
                    structured_result = {
                        'index': img_data['index'],
                        'filename': img_data['filename'],
                        'size': img_data['size'],
                        'ocr_text': ocr_result.get('text', ''),
                        'confidence': ocr_result.get('confidence', 0),
                        'word_count': ocr_result.get('word_count', 0),
                        'saved_path': image_path,
                        'preprocessing_applied': ocr_result.get('preprocessing_applied', False),
                        'best_method': ocr_result.get('method', 'unknown'),
                        'total_attempts': ocr_result.get('total_attempts', 1),
                        'ocr_score': ocr_result.get('best_score', 0)
                    }
                    
                    if 'error' in ocr_result:
                        structured_result['error'] = ocr_result['error']
                    if 'all_methods' in ocr_result:
                        structured_result['methods_tried'] = ocr_result['all_methods']
                    
                    image_ocr_results.append(structured_result)
                    
                    # Log results with advanced OCR status
                    best_method = ocr_result.get('method', 'unknown')
                    total_attempts = ocr_result.get('total_attempts', 1)
                    ocr_score = ocr_result.get('best_score', 0)
                    
                    if ocr_result.get('text', '').strip():
                        self.logger.info(f"✅ EasyOCR extracted {len(ocr_result['text'])} characters from image {img_data['index']} "
                                       f"(confidence: {ocr_result.get('confidence', 0):.2f}, words: {ocr_result.get('word_count', 0)}, "
                                       f"best method: {best_method}, attempts: {total_attempts}, score: {ocr_score:.1f})")
                    else:
                        self.logger.warning(f"❌ No text found in image {img_data['index']} "
                                          f"(tried {total_attempts} methods, best: {best_method})")
                        
                except Exception as e:
                    self.logger.error(f"EasyOCR failed for image {img_data['index']}: {e}")
                    img_data["ocr_result"] = {'text': '', 'confidence': 0, 'word_count': 0, 'error': str(e)}
                    img_data["ocr_text"] = ""
                    img_data["ocr_confidence"] = 0
                    img_data["ocr_word_count"] = 0
                    
                    image_ocr_results.append({
                        'index': img_data['index'],
                        'filename': img_data['filename'],
                        'size': img_data['size'],
                        'ocr_text': '',
                        'confidence': 0,
                        'word_count': 0,
                        'error': str(e)
                    })
            
            results['image_ocr_text'] = image_ocr_results
            
            # Create ordered content structure similar to PDFTextExtractor
            ordered_content = self._analyze_document_flow(results)
            results['ordered_content'] = ordered_content
            
            # Combine all extracted text using document flow analysis
            combined_text = self._combine_extracted_text(results)
            results['combined_text'] = combined_text
            
            # Add extraction summary similar to PDFTextExtractor
            results['extraction_summary'] = {
                'total_content_blocks': len(ordered_content),
                'text_blocks': len([b for b in ordered_content if b['type'] == 'paragraph']),
                'image_blocks': len([b for b in ordered_content if b['type'] == 'image']),
                'table_blocks': len([b for b in ordered_content if b['type'] == 'table']),
                'header_blocks': len([b for b in ordered_content if b['type'] == 'header']),
                'footer_blocks': len([b for b in ordered_content if b['type'] == 'footer']),
                'total_characters': len(combined_text),
                'processing_engine': 'EasyOCR Advanced Pipeline',
                'document_type': 'DOCX'
            }
            
            self.logger.info(f"DOCX processing completed. Total characters extracted: {len(combined_text)}, "
                           f"content blocks: {len(ordered_content)}")
            
        except Exception as e:
            self.logger.error(f"Error processing DOCX: {e}")
            raise
        
        return results
    
    def _analyze_document_flow(self, results: Dict[str, any]) -> List[Dict]:
        """
        Analyze document structure to create a reading-order flow similar to PDF position ordering.
        Since DOCX doesn't have precise coordinates, we use document structure and heuristics.
        """
        try:
            content_blocks = []
            
            # Add document elements in structural order with flow analysis
            native_text = results.get('native_text', {})
            
            # 1. Headers (top priority in document flow)
            if native_text.get('headers'):
                for i, header in enumerate(native_text['headers']):
                    content_blocks.append({
                        'type': 'header',
                        'content': header,
                        'order_priority': 100 + i,  # High priority
                        'section': 'document_header'
                    })
            
            # 2. Main paragraphs with structural analysis
            if native_text.get('paragraphs'):
                for para in native_text['paragraphs']:
                    # Analyze paragraph characteristics for better ordering
                    text = para['text']
                    para_analysis = self._analyze_paragraph_structure(text)
                    
                    content_blocks.append({
                        'type': 'paragraph',
                        'content': para['text'],
                        'order_priority': 200 + para['paragraph_num'],
                        'section': 'main_content',
                        'analysis': para_analysis
                    })
            
            # 3. Tables (interspersed with content based on position hints)
            if native_text.get('tables'):
                for table in native_text['tables']:
                    # Estimate table position based on table number and content
                    estimated_position = 250 + (table['table_num'] * 10)
                    
                    content_blocks.append({
                        'type': 'table',
                        'content': table,
                        'order_priority': estimated_position,
                        'section': 'structured_data'
                    })
            
            # 4. Images (estimated positioning based on document flow)
            if results.get('image_ocr_text'):
                for img_ocr in results['image_ocr_text']:
                    # Estimate image position (images often appear after related text)
                    estimated_position = 300 + (img_ocr['index'] * 15)
                    
                    content_blocks.append({
                        'type': 'image',
                        'content': img_ocr,
                        'order_priority': estimated_position,
                        'section': 'visual_content'
                    })
            
            # 5. Footers (bottom priority)
            if native_text.get('footers'):
                for i, footer in enumerate(native_text['footers']):
                    content_blocks.append({
                        'type': 'footer',
                        'content': footer,
                        'order_priority': 500 + i,  # Low priority
                        'section': 'document_footer'
                    })
            
            # Sort by order priority (simulates top-to-bottom reading order)
            content_blocks.sort(key=lambda x: x['order_priority'])
            
            self.logger.debug(f"Analyzed document flow: {len(content_blocks)} content blocks in reading order")
            return content_blocks
            
        except Exception as e:
            self.logger.error(f"Document flow analysis failed: {e}")
            return []
    
    def _analyze_paragraph_structure(self, text: str) -> Dict:
        """
        Analyze paragraph structure to improve document flow understanding.
        """
        try:
            analysis = {
                'length': len(text),
                'word_count': len(text.split()),
                'is_heading': False,
                'is_list_item': False,
                'has_numbers': False,
                'text_type': 'paragraph'
            }
            
            # Check if it's likely a heading (short, no punctuation at end)
            if analysis['word_count'] <= 10 and not text.rstrip().endswith(('.', '!', '?')):
                analysis['is_heading'] = True
                analysis['text_type'] = 'heading'
            
            # Check for list items
            if text.strip().startswith(('•', '-', '*', '1.', '2.', '3.')) or \
               text.strip().startswith(tuple(f'{i}.' for i in range(1, 21))):
                analysis['is_list_item'] = True
                analysis['text_type'] = 'list_item'
            
            # Check for numeric content
            import re
            if re.search(r'\d+', text):
                analysis['has_numbers'] = True
            
            return analysis
            
        except Exception as e:
            self.logger.debug(f"Paragraph analysis failed: {e}")
            return {'text_type': 'paragraph', 'length': len(text)}

    def _combine_extracted_text(self, results: Dict[str, any]) -> str:
        """
        Combine all extracted text using document flow analysis for better reading order.
        Uses structural ordering similar to position-based ordering in PDFs.
        """
        try:
            # Use document flow analysis for better ordering
            content_blocks = self._analyze_document_flow(results)
            
            if not content_blocks:
                # Fallback to original method if flow analysis fails
                return self._combine_extracted_text_fallback(results)
            
            combined_parts = []
            current_section = None
            
            for block in content_blocks:
                block_type = block['type']
                content = block['content']
                section = block['section']
                
                # Add section headers when section changes
                if section != current_section:
                    section_titles = {
                        'document_header': '=== DOCUMENT HEADERS ===',
                        'main_content': '=== MAIN CONTENT (Reading Order) ===',
                        'structured_data': '=== TABLES & STRUCTURED DATA ===',
                        'visual_content': '=== VISUAL CONTENT & IMAGES ===',
                        'document_footer': '=== DOCUMENT FOOTERS ==='
                    }
                    
                    if section in section_titles:
                        combined_parts.append(section_titles[section])
                        combined_parts.append("")
                    current_section = section
                
                # Add content based on type
                if block_type == 'header':
                    combined_parts.append(content)
                    
                elif block_type == 'paragraph':
                    analysis = block.get('analysis', {})
                    if analysis.get('is_heading'):
                        combined_parts.append(f"\n[HEADING] {content}")
                    elif analysis.get('is_list_item'):
                        combined_parts.append(f"  {content}")  # Indent list items
                    else:
                        combined_parts.append(content)
                    
                elif block_type == 'table':
                    table = content
                    combined_parts.append(f"\n[TABLE {table['table_num']}] ({table['rows']}x{table['cols']})")
                    for row in table['data']:
                        combined_parts.append(" | ".join(row))
                    combined_parts.append("")
                    
                elif block_type == 'image':
                    img_ocr = content
                    if img_ocr['ocr_text'].strip():
                        best_method = img_ocr.get('best_method', 'unknown')
                        total_attempts = img_ocr.get('total_attempts', 1)
                        ocr_score = img_ocr.get('ocr_score', 0)
                        
                        combined_parts.append(f"\n[IMAGE {img_ocr['index']}] {img_ocr['filename']}")
                        combined_parts.append(f"Confidence: {img_ocr.get('confidence', 0):.2f} | Words: {img_ocr.get('word_count', 0)} | "
                                            f"Best Method: {best_method} | Attempts: {total_attempts} | Score: {ocr_score:.1f}")
                        
                        if 'methods_tried' in img_ocr and len(img_ocr['methods_tried']) > 1:
                            methods_list = ', '.join(img_ocr['methods_tried'][:3])
                            if len(img_ocr['methods_tried']) > 3:
                                methods_list += f" (+{len(img_ocr['methods_tried']) - 3} more)"
                            combined_parts.append(f"Methods Tried: {methods_list}")
                        
                        combined_parts.append(img_ocr['ocr_text'])
                        combined_parts.append("")
                        
                elif block_type == 'footer':
                    combined_parts.append(content)
                
                # Add spacing between content blocks
                if block_type in ['paragraph', 'header', 'footer']:
                    combined_parts.append("")
            
            return '\n'.join(combined_parts)
            
        except Exception as e:
            self.logger.error(f"Document flow combination failed: {e}")
            return self._combine_extracted_text_fallback(results)
    
    def _combine_extracted_text_fallback(self, results: Dict[str, any]) -> str:
        """Fallback method for combining text using the original approach."""
        combined_parts = []
        
        native_text = results.get('native_text', {})
        
        # Add headers
        if native_text.get('headers'):
            combined_parts.append("=== HEADERS ===")
            combined_parts.extend(native_text['headers'])
            combined_parts.append("")
        
        # Add main paragraphs
        if native_text.get('paragraphs'):
            combined_parts.append("=== MAIN CONTENT ===")
            for para in native_text['paragraphs']:
                combined_parts.append(para['text'])
            combined_parts.append("")
        
        # Add tables
        if native_text.get('tables'):
            combined_parts.append("=== TABLES ===")
            for table in native_text['tables']:
                combined_parts.append(f"Table {table['table_num']} ({table['rows']}x{table['cols']}):")
                for row in table['data']:
                    combined_parts.append(" | ".join(row))
                combined_parts.append("")
        
        # Add image OCR text
        if results.get('image_ocr_text'):
            combined_parts.append("=== IMAGE OCR TEXT (Advanced EasyOCR Pipeline) ===")
            for img_ocr in results['image_ocr_text']:
                if img_ocr['ocr_text'].strip():
                    best_method = img_ocr.get('best_method', 'unknown')
                    total_attempts = img_ocr.get('total_attempts', 1)
                    ocr_score = img_ocr.get('ocr_score', 0)
                    
                    combined_parts.append(f"[IMAGE {img_ocr['index']}] {img_ocr['filename']}")
                    combined_parts.append(f"Confidence: {img_ocr.get('confidence', 0):.2f} | Words: {img_ocr.get('word_count', 0)} | "
                                        f"Best Method: {best_method} | Attempts: {total_attempts} | Score: {ocr_score:.1f}")
                    
                    if 'methods_tried' in img_ocr and len(img_ocr['methods_tried']) > 1:
                        methods_list = ', '.join(img_ocr['methods_tried'][:3])
                        if len(img_ocr['methods_tried']) > 3:
                            methods_list += f" (+{len(img_ocr['methods_tried']) - 3} more)"
                        combined_parts.append(f"Methods Tried: {methods_list}")
                    
                    combined_parts.append(img_ocr['ocr_text'])
                    combined_parts.append("")
        
        # Add footers
        if native_text.get('footers'):
            combined_parts.append("=== FOOTERS ===")
            combined_parts.extend(native_text['footers'])
        
        return '\n'.join(combined_parts)
    
    def save_extracted_text(self, results: Dict[str, any], output_filename: Optional[str] = None) -> str:
        """Save the extracted text to a file."""
        if not output_filename:
            docx_name = Path(results['docx_path']).stem
            
            # Sanitize filename to remove invalid characters for Windows
            import re
            # Remove or replace invalid characters for Windows filenames
            sanitized_name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', '_', docx_name)
            # Replace multiple underscores with single underscore
            sanitized_name = re.sub(r'_+', '_', sanitized_name)
            # Remove leading/trailing underscores
            sanitized_name = sanitized_name.strip('_')
            
            output_filename = f"{sanitized_name}_extracted_text.txt"
        
        output_path = self.output_dir / output_filename
        
        # Create summary header with EasyOCR metrics
        import datetime
        
        # Calculate advanced OCR statistics
        image_ocr_data = results.get('image_ocr_text', [])
        total_ocr_chars = sum(len(img.get('ocr_text', '')) for img in image_ocr_data)
        total_words = sum(img.get('word_count', 0) for img in image_ocr_data)
        total_attempts = sum(img.get('total_attempts', 1) for img in image_ocr_data)
        avg_confidence = 0
        avg_score = 0
        
        if image_ocr_data:
            confidences = [img.get('confidence', 0) for img in image_ocr_data if img.get('confidence', 0) > 0]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            
            scores = [img.get('ocr_score', 0) for img in image_ocr_data if img.get('ocr_score', 0) > 0]
            avg_score = sum(scores) / len(scores) if scores else 0
        
        # Count unique methods used
        all_methods = set()
        for img in image_ocr_data:
            if 'methods_tried' in img:
                all_methods.update(img['methods_tried'])
        methods_count = len(all_methods)
        
        # Enhanced summary with standardized format similar to PDFTextExtractor
        complexity_score = results.get('diagnosis', {}).get('summary', {}).get('document_complexity_score', 0)
        file_size_mb = os.path.getsize(results['docx_path']) / (1024 * 1024) if os.path.exists(results['docx_path']) else 0
        
        summary = [
            f"DOCUMENT TEXT EXTRACTION REPORT",
            f"Document Type: DOCX (Microsoft Word)",
            f"Source File: {results['docx_path']}",
            f"File Size: {file_size_mb:.2f} MB",
            f"Document Complexity Score: {complexity_score:.1f}/100",
            f"Processing Engine: EasyOCR Advanced Pipeline",
            "="*80,
            "",
            "CONTENT ANALYSIS:",
            f"  • Total Paragraphs: {results['native_text'].get('total_paragraphs', 0)}",
            f"  • Total Tables: {results['native_text'].get('total_tables', 0)}",
            f"  • Total Sections: {results.get('diagnosis', {}).get('total_sections', 0)}",
            f"  • Images Found: {len(results['images'])}",
            f"  • Total Characters: {len(results['combined_text'])}",
            "",
            "OCR PERFORMANCE METRICS:",
            f"  • OCR Characters Extracted: {total_ocr_chars}",
            f"  • OCR Words Detected: {total_words}",
            f"  • Total OCR Attempts: {total_attempts}",
            f"  • Preprocessing Methods Used: {methods_count}",
            f"  • Average OCR Confidence: {avg_confidence:.2f}",
            f"  • Average OCR Score: {avg_score:.1f}",
            "",
            "PROCESSING DETAILS:",
            f"  • Extraction Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"  • Document Flow Analysis: {'Applied' if len(results.get('ordered_content', [])) > 0 else 'Standard'}",
            f"  • Advanced Preprocessing: Enabled",
            f"  • Fallback Methods: {len(all_methods)} techniques available",
            "="*80,
            ""
        ]
        
        # Write to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(summary))
            f.write(results['combined_text'])
        
        self.logger.info(f"Extracted text saved to: {output_path}")
        return str(output_path)


# --- EasyOCR Extractor for Standalone Images ---
class EasyOCRExtractor:
    def __init__(self, enable_gpu: bool = True):
        try:
            import easyocr
            # Use the correct parameter name for GPU setting
            if enable_gpu:
                self.reader = easyocr.Reader(['en'], gpu=True)
            else:
                self.reader = easyocr.Reader(['en'], gpu=False)
        except ImportError:
            raise ImportError("EasyOCR not installed. Install with: pip install easyocr")
        except Exception as e:
            # Fallback to CPU if GPU initialization fails
            print(f"GPU initialization failed, falling back to CPU: {e}")
            self.reader = easyocr.Reader(['en'], gpu=False)
    def preprocess_image(self, image_path: str) -> np.ndarray:
        image = cv2.imread(image_path)
        if image is None:
            raise ValueError(f"Could not load image: {image_path}")
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(gray)
        height, width = enhanced.shape[:2]
        if width < 500 or height < 500:
            scale = max(500 / width, 500 / height)
            enhanced = cv2.resize(enhanced, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)
        return enhanced
    def extract_text(self, image_path: str) -> Dict:
        if not os.path.exists(image_path):
            raise FileNotFoundError(f"Image not found: {image_path}")
        try:
            results = self.reader.readtext(image_path)
            texts = []
            confidences = []
            for (bbox, text, confidence) in results:
                if confidence > 0.1:
                    texts.append(text.strip())
                    confidences.append(confidence)
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            return {
                'text': '\n'.join(texts),
                'confidence': avg_confidence,
                'word_count': len(texts)
            }
        except Exception as e:
            return {'text': '', 'confidence': 0, 'error': str(e)}

# --- PDFTextExtractor class (PyMuPDF-based) ---
class PDFTextExtractor:
    """
    A robust PDF text extraction pipeline that combines native text extraction
    with OCR-based image text detection.
    """
    
    def __init__(self, tesseract_path: Optional[str] = None, output_dir: str = "extracted_texts"):
        """
        Initialize the PDF text extractor.
        
        Args:
            tesseract_path: Path to tesseract executable (if not in PATH)
            output_dir: Directory to save extracted text files
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # Configure tesseract path if provided
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        
        # Setup logging
        import logging
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger(__name__)
    
    def diagnose_pdf(self, pdf_path: str) -> Dict[str, any]:
        """
        Diagnose PDF structure to understand content types.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Dictionary with diagnostic information
        """
        try:
            import fitz  # PyMuPDF
            pdf_document = fitz.open(pdf_path)
            diagnosis = {
                'total_pages': len(pdf_document),
                'page_details': []
            }
            
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                
                # Get basic page info
                page_info = {
                    'page_num': page_num + 1,
                    'mediabox': page.mediabox,
                    'rotation': page.rotation,
                    'image_count': len(page.get_images(full=True)),
                    'text_length': len(page.get_text()),
                    'drawings_count': 0,  # Initialize to 0
                    'links_count': len(page.get_links()),
                    'annotations_count': 0  # Initialize to 0
                }
                
                # Safely get drawings count
                try:
                    drawings = page.get_drawings()
                    page_info['drawings_count'] = len(list(drawings))  # Convert generator to list
                except:
                    page_info['drawings_count'] = 0
                
                # Safely get annotations count
                try:
                    annotations = page.annots()
                    page_info['annotations_count'] = len(list(annotations))  # Convert generator to list
                except:
                    page_info['annotations_count'] = 0
                
                # Check for different content types
                images = page.get_images(full=True)
                page_info['image_details'] = []
                
                for img_idx, img in enumerate(images):
                    try:
                        xref = img[0]
                        base_image = pdf_document.extract_image(xref)
                        img_detail = {
                            'index': img_idx + 1,
                            'xref': xref,
                            'ext': base_image.get('ext', 'unknown'),
                            'width': base_image.get('width', 0),
                            'height': base_image.get('height', 0),
                            'colorspace': base_image.get('colorspace', 'unknown'),
                            'size_bytes': len(base_image.get('image', b''))
                        }
                        page_info['image_details'].append(img_detail)
                    except Exception as e:
                        page_info['image_details'].append({
                            'index': img_idx + 1,
                            'error': str(e)
                        })
                
                diagnosis['page_details'].append(page_info)
                
                self.logger.info(f"Page {page_num + 1}: {page_info['text_length']} chars text, "
                               f"{page_info['image_count']} images, {page_info['drawings_count']} drawings")
            
            pdf_document.close()
            return diagnosis
            
        except Exception as e:
            self.logger.error(f"PDF diagnosis failed: {e}")
            return {'error': str(e)}
    
    def extract_native_text(self, pdf_document: fitz.Document) -> List[Dict]:
        """
        Extract native text from PDF pages using PyMuPDF with positional information.
        
        Args:
            pdf_document: PyMuPDF document object
            
        Returns:
            List of dictionaries containing text content and position for each page
        """
        page_texts = []
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            
            # Extract text with position information
            text_dict = page.get_text("dict")
            
            # Also get simple text for fallback
            simple_text = page.get_text("text")
            
            # Process text blocks to get position information
            text_blocks = []
            
            for block in text_dict.get("blocks", []):
                if "lines" in block:  # Text block
                    block_text = ""
                    for line in block["lines"]:
                        for span in line["spans"]:
                            block_text += span["text"]
                        block_text += "\n"
                    
                    if block_text.strip():
                        text_blocks.append({
                            "text": block_text.strip(),
                            "bbox": block["bbox"],  # (x0, y0, x1, y1)
                            "type": "text"
                        })
            
            page_data = {
                "page_num": page_num,
                "simple_text": simple_text,
                "text_blocks": text_blocks,
                "page_height": page.rect.height,
                "page_width": page.rect.width
            }
            
            page_texts.append(page_data)
            self.logger.info(f"Extracted native text from page {page_num + 1} with {len(text_blocks)} text blocks")
        
        return page_texts
    
    def extract_images_from_pdf(self, pdf_document: fitz.Document) -> List[Tuple[int, List[Dict]]]:
        """
        Enhanced image extraction from PDF pages with position information.
        
        Args:
            pdf_document: PyMuPDF document object
            
        Returns:
            List of tuples containing (page_number, list_of_image_dicts_with_position)
        """
        all_images = []
        
        # Create extracted_images folder
        extracted_images_dir = self.output_dir / "extracted_images"
        extracted_images_dir.mkdir(exist_ok=True)
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            page_images = []
            
            # Method 1: Extract images using get_images()
            image_list = page.get_images(full=True)
            self.logger.info(f"Page {page_num + 1}: Found {len(image_list)} images using get_images()")
            
            for img_index, img in enumerate(image_list):
                try:
                    # Get image data
                    xref = img[0]
                    base_image = pdf_document.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # Filter out tiny images (likely decorative elements)
                    if len(image_bytes) < 1000:  # Skip very small images
                        self.logger.debug(f"Skipping tiny image {img_index + 1} on page {page_num + 1} (size: {len(image_bytes)} bytes)")
                        continue
                    
                    # Convert to PIL Image
                    image = Image.open(io.BytesIO(image_bytes))
                    
                    # Filter out very small images by dimensions
                    if image.width < 50 or image.height < 50:
                        self.logger.debug(f"Skipping small image {img_index + 1} on page {page_num + 1} (dimensions: {image.width}x{image.height})")
                        continue
                    
                    # Save the image to disk
                    img_filename = f"page{page_num + 1}_img{img_index + 1}.png"
                    img_path = extracted_images_dir / img_filename
                    image.save(img_path)
                    
                    # Get image position on the page
                    img_bbox = None
                    try:
                        # Find image position using transformation matrix
                        for item in page.get_images(full=True):
                            if item[0] == xref:
                                # Get image rect
                                img_rects = page.get_image_rects(item)
                                if img_rects:
                                    img_bbox = img_rects[0]  # Use first rect if multiple
                                break
                    except:
                        # Fallback: estimate position (top of page)
                        img_bbox = fitz.Rect(0, 0, image.width, image.height)
                    
                    image_data = {
                        "image": image,
                        "bbox": img_bbox,
                        "index": img_index + 1,
                        "size": (image.width, image.height),
                        "type": "image",
                        "saved_path": str(img_path),
                        "filename": img_filename
                    }
                    
                    page_images.append(image_data)
                    self.logger.info(f"Successfully extracted and saved image {img_index + 1} from page {page_num + 1} (size: {image.width}x{image.height}, {len(image_bytes)} bytes) -> {img_filename}")
                    
                except Exception as e:
                    self.logger.warning(f"Failed to extract image {img_index + 1} from page {page_num + 1}: {e}")
            
            all_images.append((page_num, page_images))
            self.logger.info(f"Page {page_num + 1}: Successfully extracted {len(page_images)} valid images")
        
        return all_images
    
    def preprocess_image_for_ocr(self, image: Image.Image) -> Image.Image:
        """
        Enhanced image preprocessing to improve OCR accuracy.
        
        Args:
            image: PIL Image object
            
        Returns:
            Preprocessed PIL Image
        """
        try:
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Convert PIL to OpenCV format
            opencv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
            
            # Convert to grayscale
            gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
            
            # Apply multiple preprocessing techniques and return the best one
            processed_images = []
            
            # Method 1: Basic preprocessing
            denoised = cv2.medianBlur(gray, 3)
            thresh1 = cv2.adaptiveThreshold(
                denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
            )
            processed_images.append(('adaptive_thresh', Image.fromarray(thresh1)))
            
            # Method 2: OTSU thresholding
            _, thresh2 = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            processed_images.append(('otsu_thresh', Image.fromarray(thresh2)))
            
            # Method 3: Simple thresholding
            _, thresh3 = cv2.threshold(gray, 127, 255, cv2.THRESH_BINARY)
            processed_images.append(('simple_thresh', Image.fromarray(thresh3)))
            
            # Method 4: Enhanced contrast
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
            enhanced = clahe.apply(gray)
            _, thresh4 = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            processed_images.append(('enhanced_contrast', Image.fromarray(thresh4)))
            
            # Method 5: Morphological operations
            kernel = np.ones((2,2), np.uint8)
            morph = cv2.morphologyEx(gray, cv2.MORPH_CLOSE, kernel)
            _, thresh5 = cv2.threshold(morph, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            processed_images.append(('morphological', Image.fromarray(thresh5)))
            
            # Return the adaptive threshold as default (usually works well)
            return processed_images[0][1]
            
        except Exception as e:
            self.logger.warning(f"Image preprocessing failed: {e}, using original image")
            return image
    
    def extract_text_from_image(self, image: Image.Image) -> str:
        """
        Extract text from image using Tesseract OCR with multiple approaches.
        
        Args:
            image: PIL Image object
            
        Returns:
            Extracted text string
        """
        try:
            # Test if Tesseract is available
            try:
                version = pytesseract.get_tesseract_version()
                self.logger.debug(f"Using Tesseract version: {version}")
            except Exception as e:
                self.logger.error(f"Tesseract not available: {e}")
                return ""
            
            all_results = []
            
            # Try 1: Original image with different PSM modes
            original_configs = [
                ('original_psm6', r'--oem 3 --psm 6'),
                ('original_psm7', r'--oem 3 --psm 7'),
                ('original_psm8', r'--oem 3 --psm 8'),
                ('original_psm11', r'--oem 3 --psm 11'),
                ('original_psm13', r'--oem 3 --psm 13')
            ]
            
            for name, config in original_configs:
                try:
                    text = pytesseract.image_to_string(image, config=config).strip()
                    if text:
                        all_results.append((name, text, len(text)))
                        self.logger.debug(f"{name}: extracted {len(text)} characters")
                except Exception as e:
                    self.logger.debug(f"{name} failed: {e}")
            
            # Try 2: Preprocessed images
            try:
                # Convert to RGB if necessary
                if image.mode != 'RGB':
                    rgb_image = image.convert('RGB')
                else:
                    rgb_image = image
                
                # Convert to numpy array
                img_array = np.array(rgb_image)
                
                # Try different preprocessing approaches
                preprocessing_methods = [
                    ('grayscale', self._to_grayscale),
                    ('threshold_otsu', self._apply_otsu_threshold),
                    ('threshold_adaptive', self._apply_adaptive_threshold),
                    ('enhance_contrast', self._enhance_contrast),
                    ('remove_noise', self._remove_noise),
                    ('resize_2x', self._resize_image),
                ]
                
                for method_name, method_func in preprocessing_methods:
                    try:
                        processed_img = method_func(img_array)
                        processed_pil = Image.fromarray(processed_img)
                        
                        # Try OCR with different configs
                        configs = [r'--oem 3 --psm 6', r'--oem 3 --psm 7', r'--oem 3 --psm 11']
                        for config in configs:
                            try:
                                text = pytesseract.image_to_string(processed_pil, config=config).strip()
                                if text:
                                    all_results.append((f"{method_name}_{config.split()[-1]}", text, len(text)))
                                    self.logger.debug(f"{method_name} with {config}: extracted {len(text)} characters")
                            except:
                                continue
                                
                    except Exception as e:
                        self.logger.debug(f"Preprocessing method {method_name} failed: {e}")
                        continue
            
            except Exception as e:
                self.logger.warning(f"Preprocessing attempts failed: {e}")
            
            # Try 3: Scale the image up (sometimes helps with small text)
            try:
                # Scale up by 2x
                width, height = image.size
                scaled_image = image.resize((width * 2, height * 2), Image.LANCZOS)
                
                configs = [r'--oem 3 --psm 6', r'--oem 3 --psm 7']
                for config in configs:
                    try:
                        text = pytesseract.image_to_string(scaled_image, config=config).strip()
                        if text:
                            all_results.append((f"scaled_{config.split()[-1]}", text, len(text)))
                            self.logger.debug(f"Scaled image with {config}: extracted {len(text)} characters")
                    except:
                        continue
                        
            except Exception as e:
                self.logger.debug(f"Image scaling failed: {e}")
            
            # Return the longest text found
            if all_results:
                best_result = max(all_results, key=lambda x: x[2])
                self.logger.info(f"Best OCR result from method '{best_result[0]}': {best_result[2]} characters")
                return best_result[1]
            else:
                self.logger.info("No text extracted from image with any method")
                return ""
                
        except Exception as e:
            self.logger.error(f"All OCR attempts failed: {e}")
            return ""
    
    def _to_grayscale(self, img_array):
        """Convert image to grayscale."""
        if len(img_array.shape) == 3:
            return cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        return img_array
    
    def _apply_otsu_threshold(self, img_array):
        """Apply OTSU thresholding."""
        gray = self._to_grayscale(img_array)
        _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        return thresh
    
    def _apply_adaptive_threshold(self, img_array):
        """Apply adaptive thresholding."""
        gray = self._to_grayscale(img_array)
        return cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
    
    def _enhance_contrast(self, img_array):
        """Enhance image contrast."""
        gray = self._to_grayscale(img_array)
        clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8,8))
        enhanced = clahe.apply(gray)
        _, thresh = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        return thresh
    
    def _remove_noise(self, img_array):
        """Remove noise from image."""
        gray = self._to_grayscale(img_array)
        denoised = cv2.medianBlur(gray, 3)
        _, thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        return thresh
    
    def _resize_image(self, img_array):
        """Resize image by 2x."""
        height, width = img_array.shape[:2]
        resized = cv2.resize(img_array, (width * 2, height * 2), interpolation=cv2.INTER_CUBIC)
        return self._to_grayscale(resized)
    
    def combine_text_and_images_by_position(self, page_text_data: Dict, page_images: List[Dict]) -> List[Dict]:
        """
        Combine text blocks and images based on their position on the page to maintain reading order.
        
        Args:
            page_text_data: Dictionary containing text blocks and page info
            page_images: List of image dictionaries with position info
            
        Returns:
            List of content blocks sorted by reading order (top to bottom, left to right)
        """
        content_blocks = []
        
        # Add text blocks
        for text_block in page_text_data.get("text_blocks", []):
            content_blocks.append({
                "type": "text",
                "content": text_block["text"],
                "bbox": text_block["bbox"],
                "y_position": text_block["bbox"][1]  # y0 coordinate for sorting
            })
        
        # Add image blocks (with OCR text if available)
        for img_data in page_images:
            bbox = img_data.get("bbox")
            if bbox:
                y_pos = bbox.y0 if hasattr(bbox, 'y0') else bbox[1] if isinstance(bbox, (list, tuple)) else 0
            else:
                y_pos = 0  # Default to top if no position info
            
            content_blocks.append({
                "type": "image",
                "content": img_data,
                "bbox": bbox,
                "y_position": y_pos
            })
        
        # Sort by vertical position (top to bottom)
        content_blocks.sort(key=lambda x: x["y_position"])
        
        return content_blocks
    
    def extract_text_from_pdf_as_images(self, pdf_path: str) -> List[str]:
        """
        Convert PDF pages to images and extract text using OCR.
        This is useful for scanned PDFs or when native text extraction fails.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            List of text content for each page
        """
        try:
            from pdf2image import convert_from_path
            
            # Convert PDF pages to images
            images = convert_from_path(pdf_path, dpi=300)
            page_texts = []
            
            for i, image in enumerate(images):
                text = self.extract_text_from_image(image)
                page_texts.append(text)
                self.logger.info(f"Extracted OCR text from PDF page {i + 1} (as image)")
            
            return page_texts
            
        except Exception as e:
            self.logger.error(f"Failed to convert PDF to images: {e}")
            return []
    
    def process_pdf(self, pdf_path: str, use_fallback_ocr: bool = True, diagnose: bool = True) -> Dict[str, any]:
        """
        Main method to process a PDF file and extract all text content.
        
        Args:
            pdf_path: Path to the PDF file
            use_fallback_ocr: Whether to use full-page OCR as fallback
            diagnose: Whether to run PDF diagnostics first
            
        Returns:
            Dictionary containing extraction results
        """
        self.logger.info(f"Starting PDF processing: {pdf_path}")
        
        results = {
            'pdf_path': pdf_path,
            'native_text': [],
            'image_ocr_text': [],
            'fallback_ocr_text': [],
            'combined_text': '',
            'total_pages': 0,
            'images_found': 0,
            'diagnosis': None
        }
        
        try:
            # Run diagnostics first if requested
            if diagnose:
                self.logger.info("Running PDF diagnostics...")
                diagnosis = self.diagnose_pdf(pdf_path)
                results['diagnosis'] = diagnosis
                
                # Log summary
                total_images = sum(page['image_count'] for page in diagnosis.get('page_details', []))
                self.logger.info(f"Diagnosis complete: {diagnosis.get('total_pages', 0)} pages, {total_images} total images detected")
            
            # Open PDF document
            pdf_document = fitz.open(pdf_path)
            results['total_pages'] = len(pdf_document)
            
            # Extract native text with position information
            self.logger.info("Extracting native text...")
            native_texts = self.extract_native_text(pdf_document)
            results['native_text'] = native_texts
            
            # Extract images and perform OCR with position information
            self.logger.info("Extracting images and performing OCR...")
            page_images = self.extract_images_from_pdf(pdf_document)
            
            # Process each page to maintain reading order
            ordered_content = []
            total_images = 0
            
            for page_num, images in page_images:
                page_image_count = len(images)
                total_images += page_image_count
                
                self.logger.info(f"Processing {page_image_count} images from page {page_num + 1}")
                
                # Process images and extract OCR text
                processed_images = []
                for img_idx, image_data in enumerate(images):
                    try:
                        image = image_data["image"]
                        self.logger.info(f"Starting OCR for image {img_idx + 1} on page {page_num + 1} (size: {image.size})")
                        
                        # Try enhanced OCR without debug output
                        ocr_text = self.extract_text_from_image(image)
                        
                        # Add OCR text to image data
                        image_data["ocr_text"] = ocr_text
                        processed_images.append(image_data)
                        
                        if ocr_text.strip():
                            self.logger.info(f"✅ Extracted {len(ocr_text)} characters from image {img_idx + 1} on page {page_num + 1}")
                        else:
                            self.logger.warning(f"❌ No text found in image {img_idx + 1} on page {page_num + 1}")
                            
                            # For debugging, save problematic image only if needed
                            if page_num == 0 and img_idx == 0:  # Only for first image to avoid spam
                                self.logger.debug("First image had no OCR results - this may indicate OCR configuration issues")
                                    
                    except Exception as e:
                        self.logger.error(f"OCR failed for image {img_idx + 1} on page {page_num + 1}: {e}")
                        image_data["ocr_text"] = ""
                        processed_images.append(image_data)
                
                # Combine text and images by position for this page
                if page_num < len(native_texts):
                    page_content = self.combine_text_and_images_by_position(
                        native_texts[page_num], processed_images
                    )
                    ordered_content.append(page_content)
                else:
                    # Page with only images
                    ordered_content.append([{
                        "type": "image",
                        "content": img_data,
                        "bbox": img_data.get("bbox"),
                        "y_position": 0
                    } for img_data in processed_images])
            
            results['ordered_content'] = ordered_content
            results['images_found'] = total_images
            
            # Close the PDF document
            pdf_document.close()
            
            # Fallback: Use full-page OCR if native text extraction yields poor results
            if use_fallback_ocr:
                total_native_chars = 0
                for page_data in native_texts:
                    if isinstance(page_data, dict):
                        total_native_chars += len(page_data.get('simple_text', ''))
                    else:
                        total_native_chars += len(str(page_data))
                
                if total_native_chars < 100:  # Threshold for "poor" native extraction
                    self.logger.info("Native text extraction yielded minimal results. Using full-page OCR...")
                    fallback_texts = self.extract_text_from_pdf_as_images(pdf_path)
                    results['fallback_ocr_text'] = fallback_texts
            
            # Combine all extracted text
            combined_text = self._combine_extracted_text(results)
            results['combined_text'] = combined_text
            
            self.logger.info(f"PDF processing completed. Total characters extracted: {len(combined_text)}")
            
        except Exception as e:
            self.logger.error(f"Error processing PDF: {e}")
            raise
        
        return results
    
    def _combine_extracted_text(self, results: Dict[str, any]) -> str:
        """
        Combine all extracted text into a single string, maintaining reading order.
        
        Args:
            results: Dictionary containing extraction results
            
        Returns:
            Combined text string with preserved reading order
        """
        combined_parts = []
        
        # Use ordered content if available (new method)
        if 'ordered_content' in results and results['ordered_content']:
            for page_num, page_content in enumerate(results['ordered_content']):
                if not page_content:
                    continue
                    
                combined_parts.append(f"\n=== PAGE {page_num + 1} ===")
                
                for block in page_content:
                    if block["type"] == "text":
                        combined_parts.append(block["content"])
                    elif block["type"] == "image":
                        img_data = block["content"]
                        ocr_text = img_data.get("ocr_text", "")
                        if ocr_text.strip():
                            combined_parts.append(f"\n[IMAGE {img_data.get('index', '?')}]")
                            combined_parts.append(ocr_text)
                        else:
                            combined_parts.append(f"\n[IMAGE {img_data.get('index', '?')} - No text detected]")
                
                combined_parts.append("\n" + "="*50 + "\n")
        
        # Fallback to old method if ordered content not available
        else:
            for page_num in range(results['total_pages']):
                page_parts = []
                
                # Add native text
                if page_num < len(results.get('native_text', [])):
                    if isinstance(results['native_text'][page_num], dict):
                        # New format with position info
                        simple_text = results['native_text'][page_num].get('simple_text', '')
                        if simple_text.strip():
                            page_parts.append("=== NATIVE TEXT ===")
                            page_parts.append(simple_text)
                    elif isinstance(results['native_text'][page_num], str):
                        # Old format
                        if results['native_text'][page_num].strip():
                            page_parts.append("=== NATIVE TEXT ===")
                            page_parts.append(results['native_text'][page_num])
                
                # Add image OCR text
                if (results.get('image_ocr_text') and 
                    page_num < len(results['image_ocr_text']) and 
                    results['image_ocr_text'][page_num].strip()):
                    page_parts.append("=== IMAGE OCR TEXT ===")
                    page_parts.append(results['image_ocr_text'][page_num])
                
                # Add fallback OCR text
                if (results.get('fallback_ocr_text') and 
                    page_num < len(results['fallback_ocr_text']) and 
                    results['fallback_ocr_text'][page_num].strip()):
                    page_parts.append("=== FALLBACK OCR TEXT ===")
                    page_parts.append(results['fallback_ocr_text'][page_num])
                
                if page_parts:
                    combined_parts.append(f"\n=== PAGE {page_num + 1} ===")
                    combined_parts.extend(page_parts)
                    combined_parts.append("\n" + "="*50 + "\n")
        
        return '\n'.join(combined_parts)
    
    def save_extracted_text(self, results: Dict[str, any], output_filename: Optional[str] = None) -> str:
        """
        Save the extracted text to a file.
        
        Args:
            results: Dictionary containing extraction results
            output_filename: Custom output filename (optional)
            
        Returns:
            Path to the saved file
        """
        if not output_filename:
            pdf_name = Path(results['pdf_path']).stem
            output_filename = f"{pdf_name}_extracted_text.txt"
        
        output_path = self.output_dir / output_filename
        
        # Create summary header
        import logging
        summary = [
            f"PDF TEXT EXTRACTION REPORT",
            f"Source PDF: {results['pdf_path']}",
            f"Total Pages: {results['total_pages']}",
            f"Images Found: {results['images_found']}",
            f"Total Characters: {len(results['combined_text'])}",
            f"Extraction Date: {logging.Formatter().formatTime(logging.LogRecord('', 0, '', 0, '', (), None))}",
            "="*80,
            ""
        ]
        
        # Write to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(summary))
            f.write(results['combined_text'])
        
        self.logger.info(f"Extracted text saved to: {output_path}")
        return str(output_path)


# --- Main function to handle both PDFs and images ---
def main():
    # Set your Tesseract path if needed
    tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    file_path = "https://res.cloudinary.com/dewqsghdi/image/upload/v1741963429/TIT-Immobilier/aj4n1h6twn3mikr9731c.pdf"  # Change this to your file (PDF or image)
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        print(f"\n[PDF Mode] Processing: {file_path}")
        extractor = PDFTextExtractor(tesseract_path=tesseract_path)
        results = extractor.process_pdf(file_path)
        print("\n=== PDF Native Text ===")
        for page in results['native_text']:
            print(f"\n--- Page {page['page_num']} ---\n{page['text']}")
        print("\n=== PDF Image OCR Results ===")
        if 'ordered_content' in results:
            for page_num, page_content in enumerate(results['ordered_content']):
                for block in page_content:
                    if block['type'] == 'image':
                        img_data = block['content']
                        filename = img_data.get('filename', 'unknown')
                        ocr_text = img_data.get('ocr_text', '')
                        print(f"Page {page_num + 1} Image {img_data['index']} (saved as: {filename}):")
                        print(f"Size: {img_data['size']}")
                        if ocr_text:
                            print(f"OCR Text: {ocr_text[:100]}{'...' if len(ocr_text) > 100 else ''}")
                        else:
                            print("OCR Text: No text detected")
                        print()
        else:
            print("No image extraction results available")
    elif ext == '.docx':
        print(f"\n[DOCX Mode] Processing: {file_path}")
        extractor = DocxTextExtractor(enable_gpu=True)  # Use EasyOCR with GPU if available
        results = extractor.process_docx(file_path)
        
        # Display standardized summary similar to PDFTextExtractor
        extraction_summary = results.get('extraction_summary', {})
        complexity_score = results.get('diagnosis', {}).get('summary', {}).get('document_complexity_score', 0)
        
        print(f"\n📄 DOCX Processing Summary:")
        print(f"  Document Type: {extraction_summary.get('document_type', 'DOCX')}")
        print(f"  Processing Engine: {extraction_summary.get('processing_engine', 'EasyOCR')}")
        print(f"  Complexity Score: {complexity_score:.1f}/100")
        print(f"  Total Content Blocks: {extraction_summary.get('total_content_blocks', 0)}")
        print(f"  Total Characters: {extraction_summary.get('total_characters', 0)}")
        
        print(f"\n📊 Content Breakdown:")
        print(f"  • Text Blocks: {extraction_summary.get('text_blocks', 0)}")
        print(f"  • Table Blocks: {extraction_summary.get('table_blocks', 0)}")
        print(f"  • Image Blocks: {extraction_summary.get('image_blocks', 0)}")
        print(f"  • Header Blocks: {extraction_summary.get('header_blocks', 0)}")
        print(f"  • Footer Blocks: {extraction_summary.get('footer_blocks', 0)}")
        
        # Display first few paragraphs
        native_text = results['native_text']
        print("\n=== DOCX Native Text (First 3 Paragraphs) ===")
        for i, para in enumerate(native_text.get('paragraphs', [])[:3]):
            print(f"Paragraph {para['paragraph_num']}: {para['text'][:200]}{'...' if len(para['text']) > 200 else ''}")
        
        # Display tables
        if native_text.get('tables'):
            print(f"\n=== DOCX Tables ({len(native_text['tables'])}) ===")
            for table in native_text['tables'][:2]:  # Show first 2 tables
                print(f"Table {table['table_num']} ({table['rows']}x{table['cols']}):")
                for row in table['data'][:3]:  # Show first 3 rows
                    print("  " + " | ".join(row))
                if table['rows'] > 3:
                    print("  ...")
        
        # Display image OCR results with advanced EasyOCR metrics
        print("\n=== DOCX Image OCR Results (Advanced EasyOCR Pipeline) ===")
        for img_ocr in results['image_ocr_text']:
            filename = img_ocr['filename']
            ocr_text = img_ocr['ocr_text']
            confidence = img_ocr.get('confidence', 0)
            word_count = img_ocr.get('word_count', 0)
            best_method = img_ocr.get('best_method', 'unknown')
            total_attempts = img_ocr.get('total_attempts', 1)
            ocr_score = img_ocr.get('ocr_score', 0)
            
            print(f"Image {img_ocr['index']} (saved as: {filename}):")
            print(f"Size: {img_ocr['size']}")
            print(f"Confidence: {confidence:.2f} | Words detected: {word_count}")
            print(f"Best Method: {best_method} | Attempts: {total_attempts} | Score: {ocr_score:.1f}")
            
            if 'methods_tried' in img_ocr:
                methods_list = ', '.join(img_ocr['methods_tried'][:5])  # Show first 5 methods
                if len(img_ocr['methods_tried']) > 5:
                    methods_list += f" (+{len(img_ocr['methods_tried']) - 5} more)"
                print(f"Methods Tried: {methods_list}")
            
            if ocr_text:
                print(f"OCR Text: {ocr_text[:100]}{'...' if len(ocr_text) > 100 else ''}")
            else:
                print("OCR Text: No text detected")
            
            if 'error' in img_ocr:
                print(f"Error: {img_ocr['error']}")
            print()
        
        # Save extracted text
        extractor.save_extracted_text(results)
        
    elif ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
        print(f"\n[Image Mode] Processing: {file_path}")
        extractor = EasyOCRExtractor(enable_gpu=False)
        
        # First, preprocess the image using the preprocess_image method
        print("🔧 Preprocessing image...")
        try:
            preprocessed_image = extractor.preprocess_image(file_path)
            
            # Save preprocessed image temporarily for OCR
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                temp_path = temp_file.name
                import cv2
                cv2.imwrite(temp_path, preprocessed_image)
            
            try:
                # Now extract text from the preprocessed image
                result = extractor.extract_text(temp_path)
                print(f"✅ Preprocessing completed successfully")
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_path)
                except:
                    pass
                    
        except Exception as e:
            print(f"⚠️ Preprocessing failed: {e}")
            print("🔄 Falling back to direct OCR without preprocessing...")
            result = extractor.extract_text(file_path)
        
        print(f"Confidence: {result['confidence']:.2f}")
        print(f"Words extracted: {result.get('word_count', 0)}")
        print("\nExtracted Text:")
        print(result['text'])
        if 'error' in result:
            print(f"Error: {result['error']}")
        
        # Save extracted text to file
        if result['text'].strip():
            output_dir = Path("extracted_texts")
            output_dir.mkdir(exist_ok=True)
            
            # Create filename based on input image
            image_name = Path(file_path).stem
            output_filename = f"{image_name}_extracted_text.txt"
            output_path = output_dir / output_filename
            
            # Create summary header
            import datetime
            summary = [
                f"IMAGE TEXT EXTRACTION REPORT",
                f"Source Image: {file_path}",
                f"Processing: EasyOCR with Preprocessing",
                f"Confidence: {result['confidence']:.2f}",
                f"Words Extracted: {result.get('word_count', 0)}",
                f"Total Characters: {len(result['text'])}",
                f"Extraction Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                "="*80,
                "",
                result['text']
            ]
            
            # Write to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(summary))
            
            print(f"\nExtracted text saved to: {output_path}")
        else:
            print("\nNo text extracted - nothing to save")
    else:
        print(f"Unsupported file type: {ext}")

if __name__ == "__main__":
    main()