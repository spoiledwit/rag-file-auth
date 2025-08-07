import os
import logging
import requests
import tempfile
import zipfile
import io
from PIL import Image
from pathlib import Path
from typing import List, Tuple, Dict, Optional, Any
from urllib.parse import urlparse
import mimetypes

# PDF processing imports
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# DOCX processing imports
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# OCR imports
try:
    import easyocr
    import cv2
    import numpy as np
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False


class UniversalTextExtractor:
    """
    Universal text extractor that handles PDFs, DOCX, and images from local files or URLs.
    Supports Cloudinary URLs and uses OCR for images and image-heavy documents.
    """
    
    def __init__(self, output_dir: str = "extracted_texts", use_ocr: bool = True, enable_gpu: bool = True):
        """
        Initialize the universal text extractor.
        
        Args:
            output_dir: Directory to save extracted text files
            use_ocr: Whether to use OCR for images and image-heavy documents
            enable_gpu: Whether to enable GPU acceleration for OCR
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        self.temp_files_dir = self.output_dir / "temp_files"
        self.temp_files_dir.mkdir(exist_ok=True)
        
        self.temp_images_dir = self.output_dir / "temp_images"
        self.temp_images_dir.mkdir(exist_ok=True)
        
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger(__name__)
        
        self.use_ocr = use_ocr and EASYOCR_AVAILABLE
        self.ocr_reader = None
        
        if self.use_ocr:
            try:
                self.logger.info("Initializing EasyOCR...")
                self.ocr_reader = easyocr.Reader(['en'], gpu=enable_gpu)
                self.logger.info(f"EasyOCR initialized successfully with {'GPU' if enable_gpu else 'CPU'}")
            except Exception as e:
                self.logger.warning(f"Failed to initialize EasyOCR: {e}")
                self.use_ocr = False
        
        # Check for required libraries
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check and log available dependencies"""
        deps = {
            "PyPDF2": PYPDF2_AVAILABLE,
            "PyMuPDF": PYMUPDF_AVAILABLE,
            "python-docx": DOCX_AVAILABLE,
            "EasyOCR": EASYOCR_AVAILABLE
        }
        
        for dep, available in deps.items():
            status = "✓" if available else "✗"
            self.logger.info(f"{status} {dep}: {'Available' if available else 'Not installed'}")
    
    def is_url(self, path: str) -> bool:
        """Check if the path is a URL"""
        try:
            result = urlparse(path)
            return all([result.scheme, result.netloc])
        except:
            return False
    
    def detect_file_type(self, file_input: str) -> str:
        """
        Detect file type from URL or local file path.
        
        Args:
            file_input: File path or URL
            
        Returns:
            File type: 'pdf', 'docx', 'image', or 'unknown'
        """
        if self.is_url(file_input):
            # For URLs, try to detect from URL path or content-type
            url_path = urlparse(file_input).path.lower()
            if url_path.endswith('.pdf'):
                return 'pdf'
            elif url_path.endswith(('.docx', '.doc')):
                return 'docx'
            elif url_path.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp')):
                return 'image'
            else:
                # For Cloudinary and other URLs, try to determine from headers
                try:
                    response = requests.head(file_input, timeout=10)
                    content_type = response.headers.get('content-type', '').lower()
                    if 'pdf' in content_type:
                        return 'pdf'
                    elif 'word' in content_type or 'documentml' in content_type:
                        return 'docx'
                    elif 'image' in content_type:
                        return 'image'
                except:
                    pass
                return 'unknown'
        else:
            # Local file - use file extension and MIME type
            file_path = Path(file_input)
            extension = file_path.suffix.lower()
            
            if extension == '.pdf':
                return 'pdf'
            elif extension in ['.docx', '.doc']:
                return 'docx'
            elif extension in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp']:
                return 'image'
            else:
                # Try MIME type detection
                mime_type, _ = mimetypes.guess_type(file_input)
                if mime_type:
                    if 'pdf' in mime_type:
                        return 'pdf'
                    elif 'word' in mime_type or 'documentml' in mime_type:
                        return 'docx'
                    elif 'image' in mime_type:
                        return 'image'
                return 'unknown'
    
    def download_file(self, url: str, file_type: str) -> str:
        """
        Download file from URL to temporary location.
        
        Args:
            url: URL to download from
            file_type: Type of file (for extension)
            
        Returns:
            Path to downloaded temporary file
        """
        try:
            self.logger.info(f"Downloading {file_type} from URL: {url}")
            response = requests.get(url, timeout=120)
            response.raise_for_status()
            
            # Determine file extension
            extensions = {
                'pdf': '.pdf',
                'docx': '.docx',
                'image': '.png'  # Default for images
            }
            
            # Try to get extension from URL or content-type
            url_path = urlparse(url).path
            if url_path and '.' in url_path:
                detected_ext = Path(url_path).suffix
                if detected_ext:
                    extension = detected_ext
                else:
                    extension = extensions.get(file_type, '.tmp')
            else:
                extension = extensions.get(file_type, '.tmp')
            
            # Create temporary file
            temp_file = tempfile.NamedTemporaryFile(
                delete=False,
                suffix=extension,
                dir=self.temp_files_dir
            )
            temp_file.write(response.content)
            temp_file.close()
            
            self.logger.info(f"Downloaded {file_type} to: {temp_file.name} ({len(response.content)} bytes)")
            return temp_file.name
            
        except requests.RequestException as e:
            self.logger.error(f"Failed to download file from URL: {e}")
            raise ValueError(f"Failed to download file from URL: {e}")
        except Exception as e:
            self.logger.error(f"Unexpected error downloading file: {e}")
            raise
    
    def get_local_file_path(self, file_input: str, file_type: str) -> Tuple[str, bool]:
        """
        Get local file path, downloading if it's a URL.
        
        Args:
            file_input: File path or URL
            file_type: Detected file type
            
        Returns:
            Tuple of (local_path, is_temporary)
        """
        if self.is_url(file_input):
            temp_path = self.download_file(file_input, file_type)
            return temp_path, True
        else:
            if not os.path.exists(file_input):
                raise FileNotFoundError(f"File not found: {file_input}")
            return file_input, False
    
    def extract_text_from_pdf(self, pdf_path: str) -> Dict[str, Any]:
        """
        Extract text from PDF using multiple approaches.
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        result = {
            'text': '',
            'pages': 0,
            'method': 'none',
            'images_found': 0,
            'images_processed': 0,
            'ocr_text': '',
            'error': None
        }
        
        if not (PYPDF2_AVAILABLE or PYMUPDF_AVAILABLE):
            result['error'] = "No PDF libraries available. Install PyPDF2 or PyMuPDF."
            return result
        
        # Try PyMuPDF first (better for complex PDFs)
        if PYMUPDF_AVAILABLE:
            try:
                self.logger.info("Attempting PDF extraction with PyMuPDF...")
                result = self._extract_pdf_with_pymupdf(pdf_path)
                self.logger.info(f"PyMuPDF extraction successful: {result['pages']} pages, {len(result.get('text', ''))} characters")
                return result
            except Exception as e:
                self.logger.warning(f"PyMuPDF extraction failed: {e}")
                # Fall back to PyPDF2 if available
                if PYPDF2_AVAILABLE:
                    self.logger.info("Falling back to PyPDF2...")
                    try:
                        result = self._extract_pdf_with_pypdf2(pdf_path)
                        self.logger.info(f"PyPDF2 extraction successful: {result['pages']} pages")
                        return result
                    except Exception as e2:
                        self.logger.error(f"PyPDF2 extraction also failed: {e2}")
                        result['error'] = f"Both PyMuPDF and PyPDF2 failed. PyMuPDF: {e}, PyPDF2: {e2}"
                        return result
                else:
                    result['error'] = f"PyMuPDF failed and PyPDF2 not available: {e}"
                    return result
        
        # If PyMuPDF not available, try PyPDF2
        elif PYPDF2_AVAILABLE:
            try:
                self.logger.info("Attempting PDF extraction with PyPDF2...")
                result = self._extract_pdf_with_pypdf2(pdf_path)
                self.logger.info(f"PyPDF2 extraction successful: {result['pages']} pages")
                return result
            except Exception as e:
                self.logger.error(f"PyPDF2 extraction failed: {e}")
                result['error'] = str(e)
                return result
        
        # This shouldn't happen given the check above, but just in case
        result['error'] = "No PDF extraction methods available"
        return result
    
    def _extract_pdf_with_pymupdf(self, pdf_path: str) -> Dict[str, Any]:
        """Extract text from PDF using PyMuPDF with image extraction"""
        doc = None
        try:
            doc = fitz.open(pdf_path)
            total_pages = len(doc)
            
            text_parts = []
            all_images = []
            ocr_texts = []
            
            for page_num in range(total_pages):
                try:
                    page = doc.load_page(page_num)
                    
                    # Extract text
                    page_text = page.get_text()
                    if page_text.strip():
                        text_parts.append(f"\n--- PAGE {page_num + 1} ---\n")
                        text_parts.append(page_text)
                    
                    # Extract images if OCR is enabled
                    if self.use_ocr:
                        try:
                            image_list = page.get_images()
                            for img_index, img in enumerate(image_list):
                                try:
                                    # Get image data safely
                                    xref = img[0]
                                    
                                    # Check if xref is valid
                                    if xref <= 0:
                                        continue
                                    
                                    # Create pixmap with error handling
                                    try:
                                        pix = fitz.Pixmap(doc, xref)
                                    except Exception as pix_error:
                                        self.logger.warning(f"Failed to create pixmap for image {img_index + 1} on page {page_num + 1}: {pix_error}")
                                        continue
                                    
                                    # Check if pixmap is valid and not too large
                                    if pix.width > 0 and pix.height > 0 and pix.width < 5000 and pix.height < 5000:
                                        if pix.n - pix.alpha < 4:  # GRAY or RGB
                                            # Save image to temp file for OCR
                                            img_filename = f"pdf_page_{page_num + 1}_img_{img_index + 1}.png"
                                            img_path = self.temp_images_dir / img_filename
                                            
                                            try:
                                                pix.save(str(img_path))
                                                
                                                # Perform OCR
                                                ocr_result = self.extract_text_with_ocr(str(img_path))
                                                if ocr_result.get('text', '').strip():
                                                    ocr_texts.append(f"\n[IMAGE FROM PAGE {page_num + 1}]")
                                                    ocr_texts.append(ocr_result['text'])
                                                    ocr_texts.append(f"[END IMAGE FROM PAGE {page_num + 1}]\n")
                                                
                                                all_images.append({
                                                    'page': page_num + 1,
                                                    'index': img_index + 1,
                                                    'path': str(img_path),
                                                    'ocr_text': ocr_result.get('text', ''),
                                                    'confidence': ocr_result.get('confidence', 0)
                                                })
                                            except Exception as save_error:
                                                self.logger.warning(f"Failed to save image {img_index + 1} on page {page_num + 1}: {save_error}")
                                    
                                    # Clean up pixmap
                                    if pix:
                                        pix = None
                                        
                                except Exception as img_error:
                                    self.logger.warning(f"Failed to process image {img_index + 1} on page {page_num + 1}: {img_error}")
                                    continue
                        except Exception as page_img_error:
                            self.logger.warning(f"Failed to get images from page {page_num + 1}: {page_img_error}")
                
                except Exception as page_error:
                    self.logger.warning(f"Failed to process page {page_num + 1}: {page_error}")
                    continue
            
            return {
                'text': ''.join(text_parts),
                'pages': total_pages,
                'method': 'PyMuPDF',
                'images_found': len(all_images),
                'images_processed': len([img for img in all_images if img['ocr_text']]),
                'ocr_text': ''.join(ocr_texts),
                'images': all_images
            }
            
        except Exception as e:
            self.logger.error(f"Failed to open or process PDF with PyMuPDF: {e}")
            raise
        finally:
            # Ensure document is closed
            if doc:
                try:
                    doc.close()
                except:
                    pass
    
    def _extract_pdf_with_pypdf2(self, pdf_path: str) -> Dict[str, Any]:
        """Extract text from PDF using PyPDF2 (text only)"""
        text_parts = []
        
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"\n--- PAGE {page_num + 1} ---\n")
                        text_parts.append(page_text)
                except Exception as e:
                    self.logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
        
        return {
            'text': ''.join(text_parts),
            'pages': len(reader.pages),
            'method': 'PyPDF2',
            'images_found': 0,
            'images_processed': 0,
            'ocr_text': '',
            'images': []
        }
    
    def extract_text_from_docx(self, docx_path: str) -> Dict[str, Any]:
        """
        Extract text from DOCX file including embedded images.
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        if not DOCX_AVAILABLE:
            return {'error': "python-docx not available. Install with: pip install python-docx"}
        
        try:
            doc = Document(docx_path)
            
            # Extract text content
            text_parts = []
            
            # Headers
            for section in doc.sections:
                if section.header:
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            text_parts.append(f"[HEADER] {para.text}")
            
            # Main content
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Tables
            for table_idx, table in enumerate(doc.tables):
                text_parts.append(f"\n[TABLE {table_idx + 1}]")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
                text_parts.append(f"[END TABLE {table_idx + 1}]\n")
            
            # Footers
            for section in doc.sections:
                if section.footer:
                    for para in section.footer.paragraphs:
                        if para.text.strip():
                            text_parts.append(f"[FOOTER] {para.text}")
            
            # Extract images and perform OCR
            images = []
            ocr_texts = []
            
            if self.use_ocr:
                images = self._extract_images_from_docx(docx_path)
                for img_data in images:
                    ocr_result = self.extract_text_with_ocr(img_data['path'])
                    img_data['ocr_text'] = ocr_result.get('text', '')
                    img_data['confidence'] = ocr_result.get('confidence', 0)
                    
                    if img_data['ocr_text'].strip():
                        ocr_texts.append(f"\n[IMAGE {img_data['index']}]")
                        ocr_texts.append(img_data['ocr_text'])
                        ocr_texts.append(f"[END IMAGE {img_data['index']}]\n")
            
            return {
                'text': '\n'.join(text_parts),
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'images_found': len(images),
                'images_processed': len([img for img in images if img.get('ocr_text', '')]),
                'ocr_text': ''.join(ocr_texts),
                'images': images,
                'method': 'python-docx'
            }
            
        except Exception as e:
            self.logger.error(f"DOCX extraction failed: {e}")
            return {'error': str(e)}
    
    def _extract_images_from_docx(self, docx_path: str) -> List[Dict]:
        """Extract images from DOCX file"""
        images = []
        
        try:
            with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                image_files = [f for f in docx_zip.namelist() 
                              if f.startswith('word/media/') and 
                              f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff'))]
                
                for img_index, img_file in enumerate(image_files):
                    try:
                        image_data = docx_zip.read(img_file)
                        
                        # Filter out tiny images
                        if len(image_data) < 1000:
                            continue
                        
                        # Save image
                        image_filename = f"docx_image_{img_index + 1}.png"
                        image_path = self.temp_images_dir / image_filename
                        
                        with open(image_path, 'wb') as f:
                            f.write(image_data)
                        
                        images.append({
                            'index': img_index + 1,
                            'path': str(image_path),
                            'filename': image_filename
                        })
                        
                    except Exception as e:
                        self.logger.warning(f"Failed to extract image {img_index + 1}: {e}")
            
        except Exception as e:
            self.logger.error(f"Failed to extract images from DOCX: {e}")
        
        return images
    
    def extract_text_with_ocr(self, image_path: str) -> Dict[str, Any]:
        """
        Extract text from image using OCR.
        
        Args:
            image_path: Path to image file
            
        Returns:
            Dictionary with OCR results
        """
        if not self.use_ocr:
            return {'text': '', 'confidence': 0, 'error': 'OCR not available'}
        
        try:
            results = self.ocr_reader.readtext(image_path)
            
            texts = []
            confidences = []
            
            for (bbox, text, confidence) in results:
                if confidence > 0.1:
                    texts.append(text.strip())
                    confidences.append(confidence)
            
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            
            return {
                'text': '\n'.join(texts),
                'confidence': avg_confidence,
                'word_count': len(texts)
            }
            
        except Exception as e:
            return {'text': '', 'confidence': 0, 'error': str(e)}
    
    def process_file(self, file_input: str) -> Dict[str, Any]:
        """
        Main method to process any file type.
        
        Args:
            file_input: Path to file or URL
            
        Returns:
            Dictionary with extraction results
        """
        self.logger.info(f"Processing file: {file_input}")
        
        # Detect file type
        file_type = self.detect_file_type(file_input)
        self.logger.info(f"Detected file type: {file_type}")
        
        if file_type == 'unknown':
            return {'error': f"Unsupported file type or unable to detect type for: {file_input}"}
        
        # Get local file path
        try:
            local_path, is_temp = self.get_local_file_path(file_input, file_type)
        except Exception as e:
            return {'error': str(e)}
        
        # Process based on file type
        try:
            if file_type == 'pdf':
                result = self.extract_text_from_pdf(local_path)
            elif file_type == 'docx':
                result = self.extract_text_from_docx(local_path)
            elif file_type == 'image':
                result = self.extract_text_with_ocr(local_path)
                result['method'] = 'OCR'
            else:
                result = {'error': f"Unsupported file type: {file_type}"}
            
            # Add metadata
            result['file_input'] = file_input
            result['file_type'] = file_type
            result['is_url'] = self.is_url(file_input)
            result['local_path'] = local_path
            
            return result
            
        except Exception as e:
            self.logger.error(f"Processing failed: {e}")
            return {'error': str(e)}
        finally:
            # Clean up temporary file
            if is_temp and os.path.exists(local_path):
                try:
                    os.unlink(local_path)
                    self.logger.info(f"Cleaned up temporary file: {local_path}")
                except Exception as e:
                    self.logger.warning(f"Failed to clean up temporary file: {e}")
    
    def save_results(self, results: Dict[str, Any], output_filename: Optional[str] = None) -> str:
        """
        Save extraction results to a text file.
        
        Args:
            results: Results dictionary from process_file
            output_filename: Custom output filename
            
        Returns:
            Path to saved file
        """
        if 'error' in results:
            raise ValueError(f"Cannot save results with error: {results['error']}")
        
        # Generate output filename
        if not output_filename:
            if results.get('is_url', False):
                url_path = urlparse(results['file_input']).path
                if url_path and '/' in url_path:
                    base_name = Path(url_path).stem
                else:
                    base_name = f"cloudinary_{results['file_type']}"
            else:
                base_name = Path(results['file_input']).stem
            
            output_filename = f"{base_name}_extracted_text.txt"
        
        output_path = self.output_dir / output_filename
        
        # Create comprehensive report
        file_type = results.get('file_type', 'unknown')
        method = results.get('method', 'unknown')
        
        # Build summary
        summary_lines = [
            "UNIVERSAL TEXT EXTRACTION REPORT",
            "=" * 50,
            f"Source: {results.get('file_input', 'unknown')}",
            f"File Type: {file_type.upper()}",
            f"Extraction Method: {method}",
            f"Source Type: {'URL' if results.get('is_url', False) else 'Local File'}",
        ]
        
        # Add type-specific metadata
        if file_type == 'pdf':
            summary_lines.extend([
                f"Pages: {results.get('pages', 0)}",
                f"Images Found: {results.get('images_found', 0)}",
                f"Images Processed with OCR: {results.get('images_processed', 0)}",
            ])
        elif file_type == 'docx':
            summary_lines.extend([
                f"Paragraphs: {results.get('paragraphs', 0)}",
                f"Tables: {results.get('tables', 0)}",
                f"Images Found: {results.get('images_found', 0)}",
                f"Images Processed with OCR: {results.get('images_processed', 0)}",
            ])
        elif file_type == 'image':
            summary_lines.extend([
                f"OCR Confidence: {results.get('confidence', 0):.2f}",
                f"Words Detected: {results.get('word_count', 0)}",
            ])
        
        # Add text length info
        main_text = results.get('text', '')
        ocr_text = results.get('ocr_text', '')
        total_chars = len(main_text) + len(ocr_text)
        
        summary_lines.extend([
            f"Main Text Characters: {len(main_text)}",
            f"OCR Text Characters: {len(ocr_text)}",
            f"Total Characters: {total_chars}",
            "=" * 50,
            ""
        ])
        
        # Write file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(summary_lines))
            
            # Write main text
            if main_text.strip():
                f.write("MAIN EXTRACTED TEXT:\n")
                f.write("-" * 30 + "\n")
                f.write(main_text)
                f.write("\n\n")
            
            # Write OCR text if available
            if ocr_text.strip():
                f.write("OCR EXTRACTED TEXT (FROM IMAGES):\n")
                f.write("-" * 40 + "\n")
                f.write(ocr_text)
                f.write("\n\n")
        
        self.logger.info(f"Results saved to: {output_path}")
        return str(output_path)
    
    def cleanup_temp_files(self):
        """Clean up temporary files and directories"""
        import shutil
        
        for temp_dir in [self.temp_files_dir, self.temp_images_dir]:
            try:
                if temp_dir.exists():
                    shutil.rmtree(temp_dir)
                    self.logger.info(f"Cleaned up temporary directory: {temp_dir}")
            except Exception as e:
                self.logger.warning(f"Failed to clean up {temp_dir}: {e}")


def process_file_universal(file_input: str, output_dir: str = "extracted_texts", 
                          use_ocr: bool = True, enable_gpu: bool = True) -> Dict[str, Any]:
    """
    Universal file processor for PDFs, DOCX, and images.
    
    Args:
        file_input: Path to file or URL (including Cloudinary URLs)
        output_dir: Directory to save extracted text files
        use_ocr: Whether to use OCR for images
        enable_gpu: Whether to enable GPU acceleration for OCR
        
    Returns:
        Dictionary containing extraction results and output file path
    """
    extractor = UniversalTextExtractor(
        output_dir=output_dir,
        use_ocr=use_ocr,
        enable_gpu=enable_gpu
    )
    
    try:
        # Process the file
        results = extractor.process_file(file_input)
        
        if 'error' in results:
            print(f"Error processing file: {results['error']}")
            return results
        
        # Save results
        output_file = extractor.save_results(results)
        results['output_file'] = output_file
        
        # Print summary
        file_type = results.get('file_type', 'unknown').upper()
        source_type = "URL" if results.get('is_url', False) else "local file"
        method = results.get('method', 'unknown')
        
        print(f"\n{file_type} text extraction from {source_type} completed!")
        print(f"Source: {results['file_input']}")
        print(f"Method: {method}")
        print(f"Results saved to: {output_file}")
        
        # Type-specific summary
        if results['file_type'] == 'pdf':
            print(f"Pages processed: {results.get('pages', 0)}")
            print(f"Images found: {results.get('images_found', 0)}")
            print(f"Images processed with OCR: {results.get('images_processed', 0)}")
        elif results['file_type'] == 'docx':
            print(f"Paragraphs: {results.get('paragraphs', 0)}")
            print(f"Tables: {results.get('tables', 0)}")
            print(f"Images found: {results.get('images_found', 0)}")
        elif results['file_type'] == 'image':
            print(f"OCR confidence: {results.get('confidence', 0):.2f}")
            print(f"Words detected: {results.get('word_count', 0)}")
        
        main_text_len = len(results.get('text', ''))
        ocr_text_len = len(results.get('ocr_text', ''))
        print(f"Total characters extracted: {main_text_len + ocr_text_len}")
        
        # Clean up
        extractor.cleanup_temp_files()
        print("Cleaned up temporary files")
        
        return results
        
    except Exception as e:
        print(f"Error processing file: {e}")
        import traceback
        traceback.print_exc()
        raise


def main():
    """
    Example usage of the universal text extraction pipeline.
    """
    # Example files (uncomment to test)
    
    # Local files
    # file_input = "doc.docx"
    # file_input = "pdf.pdf"
    # file_input = "ticket.jpg"
    
    # Cloudinary URL examples
    # image
    # file_input = "https://res.cloudinary.com/dewqsghdi/image/upload/v1754390172/submitted_files/1754390171_1e016219_id.jpg.jpg"
    #pdf
    # file_input = "https://res.cloudinary.com/dewqsghdi/raw/upload/v1754385908/submitted_files/1234_pdf.pdf"
    #docx
    file_input = "https://res.cloudinary.com/dewqsghdi/raw/upload/v1754476510/submitted_files/1754476508_ff7e3269_doc.docx"

    try:
        results = process_file_universal(
            file_input=file_input,
            output_dir="extracted_texts",
            use_ocr=True,
            enable_gpu=True
        )
        
        if 'error' not in results:
            print("\nProcessing completed successfully!")
        else:
            print(f"\nProcessing failed: {results['error']}")
            
    except Exception as e:
        print(f"Processing failed: {e}")


if __name__ == "__main__":
    main()