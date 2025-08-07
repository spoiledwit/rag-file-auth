from docx import Document
import zipfile
import os
import logging
from PIL import Image
from pathlib import Path
from typing import List, Tuple, Dict, Optional, Any
import io

# Import EasyOCR extractor from image.py
from image import EasyOCRExtractor


class DocxTextExtractor:
    """
    A robust DOCX text extraction pipeline that combines native text extraction
    with OCR-based image text detection, similar to PDFTextExtractor.
    """
    
    def __init__(self, output_dir: str = "extracted_texts", use_easyocr: bool = True):
        """
        Initialize the DOCX text extractor.
        
        Args:
            output_dir: Directory to save extracted text files
            use_easyocr: Whether to use EasyOCR via imported image.py class
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        self.temp_images_dir = self.output_dir / "temp_images"
        self.temp_images_dir.mkdir(exist_ok=True)
        
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger(__name__)
        
        self.use_easyocr = use_easyocr
        self.easyocr_extractor = None
        
        if self.use_easyocr:
            try:
                self.logger.info("Initializing EasyOCR from image.py...")
                self.easyocr_extractor = EasyOCRExtractor(enable_gpu=True)
                self.logger.info("EasyOCR initialized successfully with GPU")
            except Exception as e:
                self.logger.warning(f"Failed to initialize EasyOCR with GPU, trying CPU: {e}")
                try:
                    self.easyocr_extractor = EasyOCRExtractor(enable_gpu=False)
                    self.logger.info("EasyOCR initialized successfully with CPU")
                except Exception as e2:
                    self.logger.error(f"Failed to initialize EasyOCR: {e2}")
                    self.use_easyocr = False
    
    def diagnose_docx(self, docx_path: str) -> Dict[str, any]:
        """
        Diagnose DOCX structure to understand content types.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Dictionary with diagnostic information
        """
        try:
            doc = Document(docx_path)
            
            # Count images in the DOCX
            image_count = 0
            with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                image_files = [f for f in docx_zip.namelist() 
                              if f.startswith('word/media/') and 
                              f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff'))]
                image_count = len(image_files)
            
            diagnosis = {
                'total_paragraphs': len(doc.paragraphs),
                'total_tables': len(doc.tables),
                'total_sections': len(doc.sections),
                'image_count': image_count,
                'text_length': sum(len(para.text) for para in doc.paragraphs),
                'non_empty_paragraphs': len([p for p in doc.paragraphs if p.text.strip()]),
                'file_size_bytes': os.path.getsize(docx_path)
            }
            
            self.logger.info(f"DOCX diagnosis complete: {diagnosis['total_paragraphs']} paragraphs, "
                           f"{diagnosis['total_tables']} tables, {diagnosis['image_count']} images")
            return diagnosis
            
        except Exception as e:
            self.logger.error(f"DOCX diagnosis failed: {e}")
            return {'error': str(e)}
    
    def extract_native_text(self, docx_path: str) -> Dict:
        """
        Extract native text from DOCX using python-docx.
        Simple approach that processes paragraphs and tables directly.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Dictionary containing structured text content
        """
        try:
            doc = Document(docx_path)
            
            # Extract paragraphs with their order
            paragraphs = []
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    paragraphs.append({
                        "element_type": "paragraph",
                        "paragraph_num": para_idx + 1,
                        "text": paragraph.text.strip(),
                        "document_order": para_idx + 1,  # Simple sequential order
                        "type": "paragraph"
                    })
            
            # Extract tables with their order  
            tables = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row_idx, row in enumerate(table.rows):
                    row_cells = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_cells.append(cell_text)
                    # Add row if it has any content
                    if row_cells:
                        table_data.append(row_cells)
                
                # Add table if it has any data
                if table_data:
                    tables.append({
                        "element_type": "table",
                        "table_num": table_idx + 1,
                        "data": table_data,
                        "rows": len(table_data),
                        "cols": max(len(row) for row in table_data) if table_data else 0,
                        "document_order": len(doc.paragraphs) + table_idx + 1,  # Place after paragraphs
                        "type": "table"
                    })
            
            # Combine all content elements
            content_elements = paragraphs + tables
            
            # Extract headers and footers (these are separate from main body)
            headers = []
            footers = []
            for section_idx, section in enumerate(doc.sections):
                if section.header:
                    header_paras = section.header.paragraphs
                    for para in header_paras:
                        if para.text.strip():
                            headers.append({
                                'text': para.text.strip(),
                                'section': section_idx + 1,
                                'type': 'header',
                                'document_order': -1000 - section_idx  # Headers come first
                            })
                
                if section.footer:
                    footer_paras = section.footer.paragraphs
                    for para in footer_paras:
                        if para.text.strip():
                            footers.append({
                                'text': para.text.strip(),
                                'section': section_idx + 1,
                                'type': 'footer',
                                'document_order': 100000 + section_idx  # Footers come last
                            })
            
            text_data = {
                'content_elements': content_elements,  # New: ordered content elements
                'paragraphs': paragraphs,
                'tables': tables,
                'headers': headers,
                'footers': footers,
                'total_paragraphs': len(paragraphs),
                'total_tables': len(tables)
            }
            
            self.logger.info(f"Extracted {len(content_elements)} content elements: "
                           f"{len(paragraphs)} paragraphs, {len(tables)} tables, "
                           f"{len(headers)} headers, {len(footers)} footers from DOCX")
            return text_data
            
        except Exception as e:
            self.logger.error(f"Failed to extract native text: {e}")
            return {'error': str(e)}
    
    def extract_images_from_docx(self, docx_path: str) -> List[Dict]:
        """
        Extract images from DOCX file and save them to temp directory.
        Attempts to determine approximate document position for images.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            List of image dictionaries with metadata and estimated document order
        """
        images = []
        
        try:
            # First, get document structure to estimate image positions
            doc = Document(docx_path)
            total_elements = len([elem for elem in doc.element.body if elem.tag.endswith(('}p', '}tbl'))])
            
            with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                image_files = [f for f in docx_zip.namelist() 
                              if f.startswith('word/media/') and 
                              f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff'))]
                
                self.logger.info(f"Found {len(image_files)} images in DOCX")
                
                for img_index, img_file in enumerate(image_files):
                    try:
                        image_data = docx_zip.read(img_file)
                        
                        # Filter out tiny images
                        if len(image_data) < 1000:
                            self.logger.debug(f"Skipping tiny image {img_index + 1} (size: {len(image_data)} bytes)")
                            continue
                        
                        # Load image to check dimensions
                        image = Image.open(io.BytesIO(image_data))
                        
                        # Filter out very small images by dimensions
                        if image.width < 50 or image.height < 50:
                            self.logger.debug(f"Skipping small image {img_index + 1} (dimensions: {image.width}x{image.height})")
                            continue
                        
                        # Save image to temporary directory
                        image_filename = f"docx_image_{img_index + 1}.png"
                        image_path = self.temp_images_dir / image_filename
                        image.save(image_path, "PNG")
                        self.logger.info(f"Saved image to: {image_path}")
                        
                        # Estimate document position for images
                        # Since we can't determine exact position, distribute images evenly through the document
                        estimated_position = (img_index + 1) * (total_elements + 1) // (len(image_files) + 1)
                        
                        image_data_dict = {
                            "image_path": str(image_path),
                            "index": img_index + 1,
                            "size": (image.width, image.height),
                            "type": "image",
                            "filename": image_filename,
                            "document_order": estimated_position,  # Estimated position in document flow
                            "estimated_position": True  # Flag to indicate this is estimated
                        }
                        
                        images.append(image_data_dict)
                        
                    except Exception as e:
                        self.logger.error(f"Failed to process image {img_index + 1}: {e}")
                        continue
            
            self.logger.info(f"Successfully extracted {len(images)} valid images from DOCX with estimated positions")
            return images
            
        except Exception as e:
            self.logger.error(f"Failed to extract images from DOCX: {e}")
            return []
    
    def extract_text_using_easyocr(self, image_path: str) -> Dict[str, any]:
        """
        Extract text from image using imported EasyOCRExtractor from image.py.
        
        Args:
            image_path: Path to the image file
            
        Returns:
            Dictionary containing OCR results
        """
        if not self.use_easyocr or self.easyocr_extractor is None:
            return {'text': '', 'confidence': 0, 'error': 'EasyOCR not available'}
        
        try:
            result = self.easyocr_extractor.extract_text(image_path)
            self.logger.info(f"EasyOCR extracted {len(result.get('text', ''))} characters from {Path(image_path).name} (confidence: {result.get('confidence', 0):.2f})")
            return result
            
        except Exception as e:
            self.logger.error(f"EasyOCR extraction failed for {image_path}: {e}")
            return {'text': '', 'confidence': 0, 'error': str(e)}
    
    def combine_content_by_order(self, native_text: Dict, images: List[Dict]) -> List[Dict]:
        """
        Combine text blocks and images based on proper document order.
        Uses the document_order field to maintain the original document structure.
        
        Args:
            native_text: Dictionary containing text blocks with document order
            images: List of image dictionaries with estimated document order
            
        Returns:
            List of content blocks sorted by true document order
        """
        content_blocks = []
        
        # Add headers (always first)
        for header in native_text.get("headers", []):
            content_blocks.append({
                "type": "header",
                "content": header["text"],
                "document_order": header["document_order"]
            })
        
        # Add content elements in their original document order
        if 'content_elements' in native_text:
            for element in native_text['content_elements']:
                if element["element_type"] == "paragraph":
                    content_blocks.append({
                        "type": "paragraph",
                        "content": element["text"],
                        "document_order": element["document_order"]
                    })
                elif element["element_type"] == "table":
                    table_text = []
                    for row in element["data"]:
                        table_text.append(" | ".join(row))
                    content_blocks.append({
                        "type": "table",
                        "content": f"[TABLE {element['table_num']}]\n" + "\n".join(table_text),
                        "document_order": element["document_order"]
                    })
        else:
            # Fallback to old method if content_elements not available
            for paragraph in native_text.get("paragraphs", []):
                content_blocks.append({
                    "type": "paragraph",
                    "content": paragraph["text"],
                    "document_order": paragraph.get("document_order", paragraph.get("order_priority", 0))
                })
            
            for table in native_text.get("tables", []):
                table_text = []
                for row in table["data"]:
                    table_text.append(" | ".join(row))
                content_blocks.append({
                    "type": "table",
                    "content": f"[TABLE {table['table_num']}]\n" + "\n".join(table_text),
                    "document_order": table.get("document_order", table.get("order_priority", 0))
                })
        
        # Add images with their estimated positions
        for img_data in images:
            content_blocks.append({
                "type": "image",
                "content": img_data,
                "document_order": img_data["document_order"]
            })
        
        # Add footers (always last)
        for footer in native_text.get("footers", []):
            content_blocks.append({
                "type": "footer",
                "content": footer["text"],
                "document_order": footer["document_order"]
            })
        
        # Sort by document order to maintain original structure
        content_blocks.sort(key=lambda x: x["document_order"])
        
        self.logger.info(f"Combined {len(content_blocks)} content blocks in document order")
        return content_blocks
    
    def process_docx(self, docx_path: str, diagnose: bool = True) -> Dict[str, any]:
        """
        Main method to process a DOCX file and extract all text content.
        
        Args:
            docx_path: Path to the DOCX file
            diagnose: Whether to run DOCX diagnostics first
            
        Returns:
            Dictionary containing extraction results
        """
        self.logger.info(f"Starting DOCX processing: {docx_path}")
        
        results = {
            'docx_path': docx_path,
            'native_text': {},
            'images': [],
            'image_ocr_text': [],
            'combined_text': '',
            'diagnosis': None
        }
        
        try:
            # Run diagnostics first if requested
            if diagnose:
                self.logger.info("Running DOCX diagnostics...")
                diagnosis = self.diagnose_docx(docx_path)
                results['diagnosis'] = diagnosis
                self.logger.info(f"Diagnosis complete: {diagnosis.get('total_paragraphs', 0)} paragraphs, {diagnosis.get('image_count', 0)} images detected")
            
            # Extract native text
            self.logger.info("Extracting native text...")
            native_text = self.extract_native_text(docx_path)
            results['native_text'] = native_text
            
            # Extract images and perform OCR
            self.logger.info("Extracting images and performing OCR...")
            images = self.extract_images_from_docx(docx_path)
            results['images'] = images
            
            # Process images with EasyOCR
            image_ocr_results = []
            for img_data in images:
                try:
                    image_path = img_data["image_path"]
                    self.logger.info(f"Starting EasyOCR for image {img_data['index']}")
                    
                    if self.use_easyocr:
                        ocr_result = self.extract_text_using_easyocr(image_path)
                        ocr_text = ocr_result.get('text', '')
                        confidence = ocr_result.get('confidence', 0)
                    else:
                        ocr_text = ""
                        confidence = 0
                    
                    # Add OCR results to image data
                    img_data["ocr_text"] = ocr_text
                    img_data["ocr_confidence"] = confidence
                    
                    if ocr_text.strip():
                        self.logger.info(f"EasyOCR extracted {len(ocr_text)} characters from image {img_data['index']} (confidence: {confidence:.2f})")
                    else:
                        self.logger.warning(f"No text found in image {img_data['index']}")
                    
                    image_ocr_results.append({
                        'index': img_data['index'],
                        'filename': img_data['filename'],
                        'size': img_data['size'],
                        'ocr_text': ocr_text,
                        'confidence': confidence
                    })
                    
                except Exception as e:
                    self.logger.error(f"EasyOCR failed for image {img_data['index']}: {e}")
                    image_ocr_results.append({
                        'index': img_data['index'],
                        'filename': img_data.get('filename', 'unknown'),
                        'size': img_data.get('size', (0, 0)),
                        'ocr_text': '',
                        'confidence': 0,
                        'error': str(e)
                    })
            
            results['image_ocr_text'] = image_ocr_results
            
            # Combine text and images by document order
            ordered_content = self.combine_content_by_order(native_text, images)
            results['ordered_content'] = ordered_content
            
            # Combine all extracted text
            combined_text = self._combine_extracted_text(results)
            results['combined_text'] = combined_text
            
            self.logger.info(f"DOCX processing completed. Total characters extracted: {len(combined_text)}")
            
        except Exception as e:
            self.logger.error(f"Error processing DOCX: {e}")
            raise
        
        return results
    
    def _combine_extracted_text(self, results: Dict[str, any]) -> str:
        """
        Combine all extracted text into a single string, maintaining proper document order.
        
        Args:
            results: Dictionary containing extraction results
            
        Returns:
            Combined text string with preserved document order
        """
        combined_parts = []
        
        # Use ordered content to maintain document structure
        if 'ordered_content' in results and results['ordered_content']:
            for block in results['ordered_content']:
                if block["type"] in ["header", "paragraph", "table", "footer"]:
                    combined_parts.append(block["content"])
                elif block["type"] == "image":
                    img_data = block["content"]
                    ocr_text = img_data.get("ocr_text", "")
                    confidence = img_data.get("ocr_confidence", 0)
                    
                    # Add image marker with position info
                    position_info = " (estimated position)" if img_data.get("estimated_position", False) else ""
                    image_index = img_data.get('index', '?')
                    
                    if ocr_text.strip():
                        combined_parts.append(f"\n[IMAGE {image_index}{position_info} - EasyOCR Confidence: {confidence:.2f}]")
                        combined_parts.append(ocr_text)
                        combined_parts.append(f"[IMAGE {image_index} END]")
                    else:
                        combined_parts.append(f"\n[IMAGE {image_index}{position_info} - No text detected by EasyOCR]")
        
        return '\n'.join(combined_parts)
    
    def save_extracted_text(self, results: Dict[str, any], output_filename: Optional[str] = None) -> str:
        """
        Save the extracted text to a file.
        
        Args:
            results: Dictionary containing extraction results
            output_filename: Custom output filename (optional)
            
        Returns:
            Path to the saved file
        """
        if not output_filename:
            docx_name = Path(results['docx_path']).stem
            output_filename = f"{docx_name}_extracted_text.txt"
        
        output_path = self.output_dir / output_filename
        
        # Create summary header
        diagnosis = results.get('diagnosis', {})
        summary = [
            f"DOCX TEXT EXTRACTION REPORT",
            f"Source DOCX: {results['docx_path']}",
            f"Total Paragraphs: {diagnosis.get('total_paragraphs', 0)}",
            f"Total Tables: {diagnosis.get('total_tables', 0)}",
            f"Images Found: {diagnosis.get('image_count', 0)}",
            f"Total Characters: {len(results['combined_text'])}",
            f"File Size: {diagnosis.get('file_size_bytes', 0)} bytes",
            "="*80,
            ""
        ]
        
        # Write to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(summary))
            f.write(results['combined_text'])
        
        self.logger.info(f"Extracted text saved to: {output_path}")
        return str(output_path)
    
    def cleanup_temp_files(self):
        """
        Clean up temporary image files.
        """
        try:
            if self.temp_images_dir.exists():
                import shutil
                shutil.rmtree(self.temp_images_dir)
                self.logger.info(f"Cleaned up temporary images directory: {self.temp_images_dir}")
        except Exception as e:
            self.logger.warning(f"Failed to clean up temporary files: {e}")


def process_docx_file(docx_path: str, output_dir: str = "extracted_texts", use_easyocr: bool = True) -> Dict[str, Any]:
    """
    Process a DOCX file and extract all text content.
    
    Args:
        docx_path: Path to the DOCX file
        output_dir: Directory to save extracted text files
        use_easyocr: Whether to use EasyOCR for image text extraction
        
    Returns:
        Dictionary containing extraction results and output file path
    """
    # Initialize the extractor with EasyOCR support (imported from image.py)
    extractor = DocxTextExtractor(
        output_dir=output_dir,
        use_easyocr=use_easyocr
    )
    
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")
    
    try:
        # Extract text from DOCX using hybrid approach (native text + EasyOCR)
        results = extractor.process_docx(docx_path, diagnose=True)
        
        # Save extracted text
        output_file = extractor.save_extracted_text(results)
        results['output_file'] = output_file
        
        print(f"DOCX text extraction completed!")
        print(f"Results saved to: {output_file}")
        print(f"Total paragraphs processed: {results['native_text'].get('total_paragraphs', 0)}")
        print(f"Images found and processed: {len(results['images'])}")
        print(f"Total characters extracted: {len(results['combined_text'])}")
        
        # Clean up temporary images
        extractor.cleanup_temp_files()
        print(f"Cleaned up temporary files")
        
        return results
        
    except Exception as e:
        print(f"Error processing DOCX: {e}")
        import traceback
        traceback.print_exc()
        raise


def main():
    """
    Example usage of the DOCX text extraction pipeline.
    """
    docx_path = "doc.docx"  # Replace with your DOCX path
    
    try:
        results = process_docx_file(docx_path)
        print("DOCX processing completed successfully!")
    except Exception as e:
        print(f"DOCX processing failed: {e}")


if __name__ == "__main__":
    main()
