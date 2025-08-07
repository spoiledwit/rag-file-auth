from docx import Document
import requests
import os
# import spacy

# API key for OCR.space
API_KEY = 'K86931413288957'

# Path to the Word document
doc_path = 'doc.docx'

# # Load SpaCy model
# try:
#     nlp = spacy.load('en_core_web_sm')  # Small English model; use 'en_core_web_lg' for better accuracy if needed
# except Exception as e:
#     print(f"Error loading SpaCy model: {e}. Please install it with: pip install spacy && python -m spacy download en_core_web_sm")
#     exit()

# Initialize the document
try:
    doc = Document(doc_path)
except Exception as e:
    print(f"Error opening document: {e}")
    exit()

# Extract text from paragraphs
print("=== Text from Paragraphs ===")
paragraph_text = [para.text for para in doc.paragraphs if para.text.strip()]
if paragraph_text:
    print('\n'.join(paragraph_text))
else:
    print("No paragraph text found.")

# Extract text from tables
print("\n=== Text from Tables ===")
table_text = []
if doc.tables:
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                table_text.append(' '.join(row_text))  # Join cells for NLP
                print(' | '.join(row_text))
else:
    print("No tables found.")

# Function to extract images and apply OCR via OCR.space
def extract_images_and_ocr(doc):
    ocr_texts = []
    print("\n=== Text from Images (via OCR.space) ===")
    image_count = 0
    for i, shape in enumerate(doc.inline_shapes):
        # Check if shape is an image (avoid FutureWarning)
        if hasattr(shape, '_inline') and shape._inline.graphic.graphicData.pic is not None:
            image_count += 1
            try:
                # Get the image data
                blip = shape._inline.graphic.graphicData.pic.blipFill.blip
                rId = blip.embed
                image_part = doc.part.related_parts[rId]
                image_data = image_part.blob

                # Send to OCR.space
                url = 'https://api.ocr.space/parse/image'
                files = {'file': ('image.png', image_data, 'image/png')}
                headers = {'apikey': API_KEY}
                try:
                    response = requests.post(url, files=files, headers=headers)
                    result = response.json()
                    if result.get('IsErroredOnProcessing', True):
                        print(f"Image {i+1}: Error - {result.get('ErrorMessage', 'Unknown error')}")
                    else:
                        ocr_text = result.get('ParsedResults', [{}])[0].get('ParsedText', '')
                        if ocr_text.strip():
                            print(f"Image {i+1} OCR Result:")
                            print(ocr_text)
                            ocr_texts.append(ocr_text)
                        else:
                            print(f"Image {i+1}: No text detected.")
                except Exception as e:
                    print(f"Error processing image {i+1} with OCR.space: {e}")
            except Exception as e:
                print(f"Error accessing image {i+1}: {e}")
    if image_count == 0:
        print("No images found.")
    return ocr_texts

# Run OCR on images
ocr_texts = extract_images_and_ocr(doc)

# # Combine all extracted text for SpaCy processing
# combined_text = '\n'.join(paragraph_text + table_text + ocr_texts)
# if not combined_text.strip():
#     print("\n=== SpaCy Analysis ===")
#     print("No text available for NLP processing.")
# else:
#     print("\n=== SpaCy Analysis ===")
#     # Process with SpaCy
#     spacy_doc = nlp(combined_text)
#     
#     # Named Entity Recognition (NER)
#     print("\nNamed Entities:")
#     if spacy_doc.ents:
#         for ent in spacy_doc.ents:
#             print(f"{ent.text} ({ent.label_})")
#     else:
#         print("No named entities found.")
#     
#     # Part-of-Speech Tagging
#     print("\nPart-of-Speech Tags:")
#     for token in spacy_doc:
#         print(f"{token.text}: {token.pos_} ({spacy.explain(token.pos_)})")
    
